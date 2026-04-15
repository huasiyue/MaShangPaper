#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
LaTeX 数学公式 → OMML (Office Math Markup Language) 转换工具

将 LaTeX 数学公式转换为 Word 原生公式格式 (OMML)，
用于通过 python-docx 将公式插入 Word 文档。

依赖: latex2mathml, lxml (python-docx 自带)
"""

from lxml import etree
import latex2mathml.converter
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── 命名空间 ──────────────────────────────────────────────
M_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def _m(tag):
    """构造 OMML 限定名"""
    return f'{{{M_NS}}}{tag}'


def _local_tag(elem):
    """获取元素的本地标签名（去掉命名空间前缀）"""
    if isinstance(elem.tag, str) and '}' in elem.tag:
        return elem.tag.split('}')[-1]
    return str(elem.tag)


# ── OMML 元素构造 ─────────────────────────────────────────

def _make_run(text, italic=False):
    """创建 OMML 数学文本运行元素 <m:r>"""
    m_r = etree.Element(_m('r'))
    m_rPr = etree.SubElement(m_r, _m('rPr'))
    m_sty = etree.SubElement(m_rPr, _m('sty'))
    m_sty.set(_m('val'), 'i' if italic else 'p')
    m_t = etree.SubElement(m_r, _m('t'))
    m_t.text = text
    return m_r


def _make_nary(nary_char, children):
    """创建 n-ary 运算符元素 (求和、积分等)"""
    nary = etree.Element(_m('nary'))
    naryPr = etree.SubElement(nary, _m('naryPr'))
    chr_elem = etree.SubElement(naryPr, _m('chr'))
    chr_elem.set(_m('val'), nary_char)
    limLoc = etree.SubElement(naryPr, _m('limLoc'))
    limLoc.set(_m('val'), 'subSup')

    sub = etree.SubElement(nary, _m('sub'))
    sup = etree.SubElement(nary, _m('sup'))
    e = etree.SubElement(nary, _m('e'))

    if len(children) > 0:
        _convert_element_or_children(children[0], sub)
    if len(children) > 1:
        _convert_element_or_children(children[1], sup)
    if len(children) > 2:
        _convert_element_or_children(children[2], e)
    return nary


# ── MathML → OMML 核心转换 ─────────────────────────────────

def _convert_children(mathml_parent, omml_parent):
    """将 MathML 元素的所有子元素转换为 OMML 并追加到父元素"""
    for child in mathml_parent:
        _convert_element(child, omml_parent)


def _convert_element_or_children(mathml_elem, omml_parent):
    """转换单个元素；若是 mrow 则直接展开其子元素"""
    if _local_tag(mathml_elem) == 'mrow':
        _convert_children(mathml_elem, omml_parent)
    else:
        _convert_element(mathml_elem, omml_parent)


def _get_children(mathml_elem):
    """获取子元素列表，自动解包仅含一个子元素的 mrow"""
    result = []
    for child in mathml_elem:
        tag = _local_tag(child)
        if tag == 'mrow' and len(list(child)) == 1:
            result.append(list(child)[0])
        else:
            result.append(child)
    return result


def _get_text(elem):
    """递归获取元素及其子元素的全部文本内容"""
    text = elem.text or ''
    for child in elem:
        text += _get_text(child)
        text += child.tail or ''
    return text.strip()


def _convert_element(mathml_elem, omml_parent):
    """将单个 MathML 元素递归转换为 OMML 并追加到父元素"""
    tag = _local_tag(mathml_elem)

    # ─── 根 / 分组元素 ───────────────────────────────
    if tag in ('math', 'mrow', 'mstyle', 'mpadded'):
        _convert_children(mathml_elem, omml_parent)

    # ─── 文本类元素 ───────────────────────────────────
    elif tag == 'mi':
        text = mathml_elem.text or ''
        omml_parent.append(_make_run(text, italic=(len(text.strip()) <= 1)))
    elif tag in ('mo', 'mn'):
        text = mathml_elem.text or ''
        omml_parent.append(_make_run(text, italic=False))
    elif tag == 'mtext':
        text = mathml_elem.text or ''
        omml_parent.append(_make_run(text, italic=False))
    elif tag == 'mspace':
        omml_parent.append(_make_run(' ', italic=False))

    # ─── 上标 / 下标 ──────────────────────────────────
    elif tag == 'msup':
        children = _get_children(mathml_elem)
        sSup = etree.SubElement(omml_parent, _m('sSup'))
        e = etree.SubElement(sSup, _m('e'))
        sup = etree.SubElement(sSup, _m('sup'))
        if len(children) > 0:
            _convert_element_or_children(children[0], e)
        if len(children) > 1:
            _convert_element_or_children(children[1], sup)

    elif tag == 'msub':
        children = _get_children(mathml_elem)
        sSub = etree.SubElement(omml_parent, _m('sSub'))
        e = etree.SubElement(sSub, _m('e'))
        sub = etree.SubElement(sSub, _m('sub'))
        if len(children) > 0:
            _convert_element_or_children(children[0], e)
        if len(children) > 1:
            _convert_element_or_children(children[1], sub)

    elif tag == 'msubsup':
        children = _get_children(mathml_elem)
        sSubSup = etree.SubElement(omml_parent, _m('sSubSup'))
        e = etree.SubElement(sSubSup, _m('e'))
        sub = etree.SubElement(sSubSup, _m('sub'))
        sup = etree.SubElement(sSubSup, _m('sup'))
        if len(children) > 0:
            _convert_element_or_children(children[0], e)
        if len(children) > 1:
            _convert_element_or_children(children[1], sub)
        if len(children) > 2:
            _convert_element_or_children(children[2], sup)

    # ─── 分数 ──────────────────────────────────────────
    elif tag == 'mfrac':
        children = _get_children(mathml_elem)
        f = etree.SubElement(omml_parent, _m('f'))
        num = etree.SubElement(f, _m('num'))
        den = etree.SubElement(f, _m('den'))
        if len(children) > 0:
            _convert_element_or_children(children[0], num)
        if len(children) > 1:
            _convert_element_or_children(children[1], den)

    # ─── 根号 ──────────────────────────────────────────
    elif tag == 'msqrt':
        rad = etree.SubElement(omml_parent, _m('rad'))
        radPr = etree.SubElement(rad, _m('radPr'))
        degHide = etree.SubElement(radPr, _m('degHide'))
        degHide.set(_m('val'), '1')
        etree.SubElement(rad, _m('deg'))
        e = etree.SubElement(rad, _m('e'))
        _convert_children(mathml_elem, e)

    elif tag == 'mroot':
        children = _get_children(mathml_elem)
        rad = etree.SubElement(omml_parent, _m('rad'))
        radPr = etree.SubElement(rad, _m('radPr'))
        degHide = etree.SubElement(radPr, _m('degHide'))
        degHide.set(_m('val'), '0')
        deg = etree.SubElement(rad, _m('deg'))
        e = etree.SubElement(rad, _m('e'))
        if len(children) > 1:
            _convert_element_or_children(children[1], deg)
        if len(children) > 0:
            _convert_element_or_children(children[0], e)

    # ─── 上方标记 (mover) ────────────────────────────
    elif tag == 'mover':
        children = _get_children(mathml_elem)
        if len(children) >= 2:
            acc_text = _get_text(children[1])
            if acc_text in ('^', '~', 'ˉ', '¯', '→', '⏜', '̂', '̃', '̇', '̈', '⃗', '̅'):
                acc = etree.SubElement(omml_parent, _m('acc'))
                accPr = etree.SubElement(acc, _m('accPr'))
                chr_el = etree.SubElement(accPr, _m('chr'))
                chr_el.set(_m('val'), acc_text)
                e = etree.SubElement(acc, _m('e'))
                _convert_element_or_children(children[0], e)
            else:
                bar = etree.SubElement(omml_parent, _m('bar'))
                barPr = etree.SubElement(bar, _m('barPr'))
                pos = etree.SubElement(barPr, _m('pos'))
                pos.set(_m('val'), 'top')
                e = etree.SubElement(bar, _m('e'))
                lim = etree.SubElement(bar, _m('lim'))
                _convert_element_or_children(children[0], e)
                _convert_element_or_children(children[1], lim)

    # ─── 下方标记 (munder) ────────────────────────────
    elif tag == 'munder':
        children = _get_children(mathml_elem)
        groupChr = etree.SubElement(omml_parent, _m('groupChr'))
        groupChrPr = etree.SubElement(groupChr, _m('groupChrPr'))
        pos = etree.SubElement(groupChrPr, _m('pos'))
        pos.set(_m('val'), 'bot')
        e = etree.SubElement(groupChr, _m('e'))
        lim = etree.SubElement(groupChr, _m('lim'))
        if len(children) > 0:
            _convert_element_or_children(children[0], e)
        if len(children) > 1:
            _convert_element_or_children(children[1], lim)

    # ─── 上下标记 (munderover) ────────────────────────
    elif tag == 'munderover':
        children = _get_children(mathml_elem)
        if len(children) >= 3:
            op_text = _get_text(children[1])
            nary_map = {
                '∑': '∑', 'Σ': '∑', '∫': '∫', '∬': '∬', '∭': '∭',
                '∏': '∏', '∐': '∐', '⋃': '⋃', '⋂': '⋂',
                '⊕': '⊕', '⊗': '⊗', '∮': '∮',
            }
            nary_char = nary_map.get(op_text)
            if nary_char:
                nary = _make_nary(nary_char, [children[1], children[2], children[0]])
                omml_parent.append(nary)
            else:
                _convert_children(mathml_elem, omml_parent)
        else:
            _convert_children(mathml_elem, omml_parent)

    # ─── 定界符 (括号) ────────────────────────────────
    elif tag == 'mfenced':
        d = etree.SubElement(omml_parent, _m('d'))
        dPr = etree.SubElement(d, _m('dPr'))
        open_ch = mathml_elem.get('open', '(')
        close_ch = mathml_elem.get('close', ')')
        begChr = etree.SubElement(dPr, _m('begChr'))
        begChr.set(_m('val'), open_ch if open_ch else ' ')
        endChr = etree.SubElement(dPr, _m('endChr'))
        endChr.set(_m('val'), close_ch if close_ch else ' ')
        for child in mathml_elem:
            e = etree.SubElement(d, _m('e'))
            _convert_element_or_children(child, e)

    # ─── 矩阵 / 表格 ─────────────────────────────────
    elif tag == 'mtable':
        m = etree.SubElement(omml_parent, _m('m'))
        for child in mathml_elem:
            if _local_tag(child) == 'mtr':
                mr = etree.SubElement(m, _m('mr'))
                for td in child:
                    if _local_tag(td) == 'mtd':
                        e = etree.SubElement(mr, _m('e'))
                        _convert_children(td, e)

    # ─── 其他未知元素：尝试递归子元素 ─────────────────
    else:
        _convert_children(mathml_elem, omml_parent)


# ── 公开 API ─────────────────────────────────────────────

def latex_to_omml(latex_str, display='block'):
    """
    将 LaTeX 数学公式转换为 OMML XML 元素。

    Args:
        latex_str: LaTeX 数学公式字符串（不含 $ 定界符）
        display: 'block' 居中独立公式 | 'inline' 行内公式

    Returns:
        lxml.etree.Element: 可直接追加到 Word 段落的 OMML 元素
    """
    mathml_str = latex2mathml.converter.convert(latex_str, display=display)
    mathml = etree.fromstring(mathml_str.encode('utf-8'))

    if display == 'block':
        oMathPara = etree.Element(_m('oMathPara'))
        oMathParaPr = etree.SubElement(oMathPara, _m('oMathParaPr'))
        jc = etree.SubElement(oMathParaPr, _m('jc'))
        jc.set(_m('val'), 'center')
        oMath = etree.SubElement(oMathPara, _m('oMath'))
        _convert_children(mathml, oMath)
        return oMathPara
    else:
        oMath = etree.Element(_m('oMath'))
        _convert_children(mathml, oMath)
        return oMath


def add_block_equation(doc, latex_str):
    """向 Word 文档中添加一个居中的块级公式段落。"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.line_spacing = 1.5
    try:
        omml_elem = latex_to_omml(latex_str, display='block')
        para._element.append(omml_elem)
    except Exception:
        run = para.add_run(latex_str)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)


def add_inline_equation(para, latex_str):
    """向已有段落中插入一个行内公式。"""
    try:
        omml_elem = latex_to_omml(latex_str, display='inline')
        para._element.append(omml_elem)
    except Exception:
        run = para.add_run(f'${latex_str}$')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
