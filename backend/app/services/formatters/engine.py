#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
论文格式化引擎（规格驱动）。
所有格式参数从 ThesisFormatSpec 读取。
"""

from __future__ import annotations

import re
from datetime import datetime
from typing import TYPE_CHECKING

from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.services.formatters.report import FormatIssueLevel, FormatReport
from app.services.formatters.spec import ThesisFormatSpec

if TYPE_CHECKING:
    from docx import Document

# 对齐方式映射
_ALIGNMENT_MAP = {
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


class ThesisFormatEngine:
    """论文格式化引擎，接收 ThesisFormatSpec 规格作用于 Word 文档。"""

    def __init__(self, spec: ThesisFormatSpec) -> None:
        self.spec = spec

    def _normalize_title_token(self, text: str) -> str:
        return re.sub(r"\s+", "", text).strip().lower()

    def _is_sdfmu_template(self) -> bool:
        return self.spec.school_id == "sdfmu"

    def _get_display_title(self, text: str) -> str:
        normalized = self._normalize_title_token(text)
        for source, display in self.spec.special_title_display_map:
            if self._normalize_title_token(source) == normalized:
                return display
        return text.strip()

    def _set_paragraph_text(self, para, text: str) -> None:
        if para.text.strip() == text.strip():
            return
        para.clear()
        para.add_run(text)

    def _apply_word_style(self, para, style_name: str) -> None:
        try:
            para.style = style_name
        except (KeyError, ValueError):
            return

    def _resolve_body_line_spacing(self):
        if self.spec.spacing.body_line_spacing_rule == "exact":
            return Pt(self.spec.spacing.body_line_spacing)
        return self.spec.spacing.body_line_spacing

    def _resolve_body_first_line_indent(self):
        return Pt(self.spec.font_sizes.body * 2)

    def _resolve_front_title_size(self):
        if self.spec.school_id == "sdfmu":
            return Pt(self.spec.font_sizes.title)
        return Pt(max(self.spec.font_sizes.abstract_label, self.spec.font_sizes.heading_1))

    def _format_front_title(self, para, font_name: str, style_name: str) -> None:
        self._apply_word_style(para, style_name)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(self.spec.spacing.paragraph_spacing)
        para.paragraph_format.space_after = Pt(self.spec.spacing.paragraph_spacing)
        para.paragraph_format.line_spacing = self.spec.spacing.heading_line_spacing
        for run in para.runs:
            run.font.name = font_name
            run.font.size = self._resolve_front_title_size()
            run.font.bold = True
            run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

    def _collect_front_text(self, doc: Document, max_paragraphs: int = 120) -> str:
        chunks = [p.text for p in doc.paragraphs[:max_paragraphs]]
        for tbl in doc.tables[:5]:
            for row in tbl.rows:
                for cell in row.cells:
                    chunks.append(cell.text)
        return "\n".join(chunks)

    # ── 主入口 ────────────────────────────────────────────

    def process_document(self, doc: Document, filename: str) -> FormatReport:
        """处理整个文档：应用格式 + 审查校验，返回 FormatReport。"""
        report = FormatReport(filename=filename)
        report.add_issue(FormatIssueLevel.INFO, "文档加载", f"成功加载文档（按{self.spec.school_name}格式处理）", "")

        self._apply_page_setup(doc, report)
        self._apply_header_footer(doc, report)

        paragraphs = doc.paragraphs
        for i, para in enumerate(paragraphs):
            self._format_paragraph(para, para_index=i, all_paragraphs=paragraphs, report=report)

        self._check_abstract(doc, report)
        self._check_cover_page(doc, report)
        self._check_declaration_page(doc, report)
        self._check_required_sections(doc, report)
        self._format_references(doc, report)
        self._check_references(doc, report)
        self._check_signature(doc, report)

        return report

    def generate_report(self, report: FormatReport) -> str:
        """生成文本格式审查报告。"""
        s = self.spec
        lines = [
            "=" * 60,
            f"{s.school_name}毕业论文格式审查报告",
            "=" * 60,
            f"文件名: {report.filename}",
            f"学校模板: {s.school_name} ({s.school_id})",
            f"检查时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            "",
            "-" * 60, "统计信息", "-" * 60,
            f"问题总数: {report.total_issues}",
            f"  - 错误: {report.error_count}",
            f"  - 警告: {report.warning_count}",
            f"  - 提示: {report.info_count}",
            "",
        ]

        for level in (FormatIssueLevel.ERROR, FormatIssueLevel.WARNING, FormatIssueLevel.INFO):
            issues = [i for i in report.issues if i.level == level]
            if issues:
                lines += ["-" * 60, f"{level.value} ({len(issues)}项)", "-" * 60]
                for idx, iss in enumerate(issues, 1):
                    lines.append(f"\n[{idx}] 位置: {iss.location}")
                    lines.append(f"    问题: {iss.description}")
                    if iss.current_value and iss.expected_value:
                        lines.append(f"    当前值: {iss.current_value}")
                        lines.append(f"    期望值: {iss.expected_value}")
                    lines.append(f"    建议: {iss.suggestion}")

        lines += ["", "=" * 60, "格式规范参考", "=" * 60]
        lines.append(self._get_format_reference())
        return "\n".join(lines)

    # ── 页面设置 ──────────────────────────────────────────

    def _apply_page_setup(self, doc: Document, report: FormatReport) -> None:
        p = self.spec.page_layout
        for section in doc.sections:
            section.top_margin = Cm(p.top_margin)
            section.bottom_margin = Cm(p.bottom_margin)
            section.left_margin = Cm(p.left_margin)
            section.right_margin = Cm(p.right_margin)
            section.gutter = Cm(p.gutter)
            section.header_distance = Cm(p.header_distance)
            section.footer_distance = Cm(p.footer_distance)
        report.add_issue(FormatIssueLevel.INFO, "页面设置", f"已应用{self.spec.school_name}页面格式", "")

    def _apply_header_footer(self, doc: Document, report: FormatReport) -> None:
        skip_header_sections = 2 if self._is_sdfmu_template() and len(doc.sections) >= 3 else 0
        skip_footer_sections = 3 if self._is_sdfmu_template() and len(doc.sections) >= 4 else skip_header_sections
        for idx, section in enumerate(doc.sections):
            # 页眉
            header = section.header
            h_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            h_para.clear()
            footer = section.footer
            f_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            f_para.clear()
            if idx < skip_header_sections:
                continue
            h_para.text = self.spec.header_text
            h_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT if self._is_sdfmu_template() else WD_ALIGN_PARAGRAPH.CENTER
            h_para.paragraph_format.line_spacing = 1.0
            h_para.paragraph_format.space_before = Pt(0)
            h_para.paragraph_format.space_after = Pt(0)
            h_para.paragraph_format.first_line_indent = Pt(0)
            h_para.paragraph_format.left_indent = Pt(0)
            h_para.paragraph_format.right_indent = Pt(0)
            for run in h_para.runs:
                run.font.name = self.spec.fonts.chinese_font
                run.font.size = Pt(self.spec.font_sizes.header_footer)
                run._element.rPr.rFonts.set(qn("w:eastAsia"), self.spec.fonts.chinese_font)
            if self._is_sdfmu_template():
                p_pr = h_para._p.get_or_add_pPr()
                for child in list(p_pr):
                    if child.tag.endswith("pBdr"):
                        p_pr.remove(child)
                if idx != skip_header_sections or skip_footer_sections == skip_header_sections:
                    p_bdr = OxmlElement("w:pBdr")
                    bottom = OxmlElement("w:bottom")
                    bottom.set(qn("w:val"), "single")
                    bottom.set(qn("w:sz"), "6")
                    bottom.set(qn("w:space"), "0")
                    bottom.set(qn("w:color"), "000000")
                    p_bdr.append(bottom)
                    p_pr.append(p_bdr)

            # 页脚（页码）
            if idx >= skip_footer_sections:
                f_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                self._add_page_number(f_para)

        report.add_issue(FormatIssueLevel.INFO, "页眉页脚", "已添加页眉和页码", "")

    def _add_page_number(self, paragraph) -> None:
        run = paragraph.add_run()
        fld1 = OxmlElement("w:fldChar"); fld1.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
        fld2 = OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"), "end")
        run._r.append(fld1); run._r.append(instr); run._r.append(fld2)
        run.font.name = self.spec.fonts.english_font
        run.font.size = Pt(self.spec.font_sizes.body)

    # ── 段落格式化 ────────────────────────────────────────

    def _format_paragraph(self, para, para_index: int, all_paragraphs: list, report: FormatReport) -> None:
        text = para.text.strip()
        display_text = self._get_display_title(text) if text else text
        if text and display_text != text and self._is_special_title(text):
            self._set_paragraph_text(para, display_text)
            text = display_text
        heading_level = self._detect_heading_level(text)

        if self._is_toc_title(text):
            self._format_front_title(para, self.spec.fonts.heading_font, "TOC Heading")
        elif self._is_abstract_title(text):
            font_name = self.spec.abstract.label_font if "摘" in text or "摘要" in text else self.spec.fonts.english_font
            self._format_front_title(para, font_name, "Title")
        elif self._is_special_title(text):
            # 特殊章节（前言/结论/致谢/附录/文献综述等）按 H1 格式处理
            self._format_heading(para, 1, style_name="Heading 1")
        elif self._is_signature_paragraph(text, para_index):
            self._format_signature(para)
        elif re.match(self.spec.caption.pattern, text):
            self._format_caption(para)
        elif self.spec.fonts.abstract_font != self.spec.fonts.chinese_font and all_paragraphs and self._is_abstract_content(para_index, all_paragraphs):
            self._format_abstract_content(para)
        elif heading_level >= 0:
            style_name = {
                0: "论文标题",
                1: "Heading 1",
                2: "Heading 2",
                3: "Heading 3",
                4: "Heading 4",
            }.get(heading_level)
            self._format_heading(para, heading_level, style_name=style_name)
        else:
            self._format_body(para)

    def _format_heading(self, para, level: int, style_name: str | None = None) -> None:
        headings = self.spec.headings
        if level < 0 or level >= len(headings):
            return
        if style_name:
            self._apply_word_style(para, style_name)
        rule = headings[level]
        para.alignment = _ALIGNMENT_MAP.get(rule.alignment, WD_ALIGN_PARAGRAPH.LEFT)
        para.paragraph_format.space_before = Pt(self.spec.spacing.paragraph_spacing)
        para.paragraph_format.space_after = Pt(self.spec.spacing.paragraph_spacing)
        para.paragraph_format.line_spacing = self.spec.spacing.heading_line_spacing
        for run in para.runs:
            run.font.name = rule.font
            run.font.size = Pt(rule.size)
            run.font.bold = rule.bold
            run._element.rPr.rFonts.set(qn("w:eastAsia"), rule.font)

    def _format_signature(self, para) -> None:
        sig = self.spec.signature
        para.alignment = _ALIGNMENT_MAP.get(sig.alignment, WD_ALIGN_PARAGRAPH.CENTER)
        para.paragraph_format.line_spacing = self._resolve_body_line_spacing()
        for run in para.runs:
            run.font.name = sig.font
            run.font.size = Pt(sig.size)
            run._element.rPr.rFonts.set(qn("w:eastAsia"), sig.font)

    def _format_abstract_content(self, para) -> None:
        self._apply_word_style(para, "Normal")
        para.paragraph_format.line_spacing = self._resolve_body_line_spacing()
        para.paragraph_format.first_line_indent = Cm(self.spec.abstract.first_line_indent)
        abstract_font = self.spec.fonts.abstract_font
        for run in para.runs:
            run.font.size = Pt(self.spec.font_sizes.body)
            run.font.name = self.spec.fonts.english_font
            if run.text and any("\u4e00" <= c <= "\u9fff" for c in run.text):
                run._element.rPr.rFonts.set(qn("w:eastAsia"), abstract_font)

    def _format_body(self, para) -> None:
        self._apply_word_style(para, "Normal")
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.line_spacing = self._resolve_body_line_spacing()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.first_line_indent = self._resolve_body_first_line_indent()
        for run in para.runs:
            run.font.size = Pt(self.spec.font_sizes.body)
            run.font.name = self.spec.fonts.english_font
            if run.text and any("\u4e00" <= c <= "\u9fff" for c in run.text):
                run._element.rPr.rFonts.set(qn("w:eastAsia"), self.spec.fonts.chinese_font)

    def _format_caption(self, para) -> None:
        self._apply_word_style(para, "Caption")
        c = self.spec.caption
        para.alignment = _ALIGNMENT_MAP.get(c.alignment, WD_ALIGN_PARAGRAPH.CENTER)
        para.paragraph_format.line_spacing = c.line_spacing
        para.paragraph_format.space_before = Pt(c.space_before)
        para.paragraph_format.space_after = Pt(c.space_after)
        for run in para.runs:
            run.font.size = Pt(c.font_size)
            run.font.name = self.spec.fonts.english_font
            run._element.rPr.rFonts.set(qn("w:eastAsia"), self.spec.fonts.chinese_font)

    # ── 检测辅助 ──────────────────────────────────────────

    def _detect_heading_level(self, text: str) -> int:
        text = text.strip()
        # 特殊章节（前言/结论/致谢/附录/文献综述等）视为 H1
        if self._is_special_title(text):
            return 1
        n = self.spec.heading_numbering
        if re.match(n.level_1, text): return 1
        if re.match(n.level_2, text): return 2
        if re.match(n.level_3, text): return 3
        if re.match(n.level_4, text): return 4
        # 论文标题启发式检测
        if len(text) < 50 and not any(c in text for c in "。，；：\u201c\u201d\u2018\u2019") and text:
            return 0
        return -1

    def _is_signature_paragraph(self, text: str, para_index: int) -> bool:
        sig = self.spec.signature
        if para_index > sig.max_position:
            return False
        text = text.strip()
        count = sum(1 for k in sig.keywords if k in text)
        if count >= 2:
            return True
        for pat in sig.prefix_patterns:
            if re.match(pat, text):
                return True
        return False

    def _is_abstract_content(self, para_index: int, paragraphs: list) -> bool:
        start_idx = -1
        for i in range(min(50, len(paragraphs))):
            t = paragraphs[i].text.strip().replace(" ", "").lower()
            for kw in self.spec.abstract.keywords:
                if kw in t and (len(t) < 15 or t.startswith(kw)):
                    start_idx = i
                    break
            if start_idx >= 0:
                break

        if start_idx == -1 or para_index <= start_idx:
            return False

        boundary_idx = len(paragraphs)
        for j in range(start_idx + 1, len(paragraphs)):
            t = paragraphs[j].text.strip().replace(" ", "").lower()
            for bk in self.spec.abstract.boundary_keywords:
                if bk in t:
                    boundary_idx = j
                    break
            if boundary_idx < len(paragraphs):
                break
        return start_idx < para_index < boundary_idx

    def _is_special_title(self, text: str) -> bool:
        text_clean = self._normalize_title_token(text)
        if len(text_clean) < 20:
            for title in self.spec.special_titles:
                normalized_title = self._normalize_title_token(title)
                if text_clean == normalized_title:
                    return True
                if normalized_title in {"附录", "appendix"} and text_clean.startswith(normalized_title):
                    return True
        return False

    def _is_toc_title(self, text: str) -> bool:
        return self._normalize_title_token(text) in {
            self._normalize_title_token("目录"),
            self._normalize_title_token("目  录"),
            self._normalize_title_token("contents"),
        }

    def _is_abstract_title(self, text: str) -> bool:
        return self._normalize_title_token(text) in {
            self._normalize_title_token("摘要"),
            self._normalize_title_token("摘  要"),
            self._normalize_title_token("abstract"),
        }

    def _find_paragraph_index(self, paragraphs: list, *candidates: str, limit: int = 120) -> int:
        normalized_candidates = {self._normalize_title_token(candidate) for candidate in candidates if candidate}
        for index, para in enumerate(paragraphs[:limit]):
            if self._normalize_title_token(para.text) in normalized_candidates:
                return index
        return -1

    def _extract_keywords(self, text: str, labels: tuple[str, ...]) -> list[str]:
        normalized_text = text.strip()
        matched_label = ""
        for label in labels:
            if self._normalize_title_token(normalized_text).startswith(self._normalize_title_token(label)):
                matched_label = label
                break
        if not matched_label:
            return []

        parts = re.split(r"[:：]", normalized_text, maxsplit=1)
        if len(parts) < 2:
            return []
        return [item.strip() for item in re.split(r"[;；,，]", parts[1]) if item.strip()]

    # ── 审查校验 ──────────────────────────────────────────

    def _check_abstract(self, doc: Document, report: FormatReport) -> None:
        paragraphs = doc.paragraphs
        chinese_index = self._find_paragraph_index(paragraphs, "摘要", "摘  要", limit=80)
        english_index = self._find_paragraph_index(paragraphs, "abstract", limit=120)
        toc_index = self._find_paragraph_index(paragraphs, "目录", "目  录", limit=120)
        inline_abstract_content = ""

        if chinese_index == -1:
            for index, para in enumerate(paragraphs[:30]):
                normalized = self._normalize_title_token(para.text)
                if normalized.startswith(self._normalize_title_token("摘要")) and normalized != self._normalize_title_token("摘要"):
                    chinese_index = index
                    inline_abstract_content = para.text.strip()
                    break

        if chinese_index == -1:
            report.add_issue(FormatIssueLevel.WARNING, "文档结构", "未检测到中文摘要", "请检查文档结构")
        else:
            title_para = paragraphs[chinese_index]
            for run in title_para.runs:
                if run.text.strip() and run.font.name not in {None, self.spec.abstract.label_font}:
                    report.add_issue(
                        FormatIssueLevel.WARNING,
                        "中文摘要",
                        f"'摘要'标题字体应为{self.spec.abstract.label_font}",
                        f"修改字体为{self.spec.abstract.label_font}",
                        run.font.name or "未设置",
                        self.spec.abstract.label_font,
                    )
                    break

            boundary_candidates = [index for index in (english_index, toc_index) if index > chinese_index]
            boundary_index = min(boundary_candidates) if boundary_candidates else len(paragraphs)
            content_parts: list[str] = []
            keyword_count = 0
            if inline_abstract_content:
                content_parts.append(re.sub(r"^\s*摘\s*要\s*[:：]?\s*", "", inline_abstract_content))
            for para in paragraphs[chinese_index + 1:boundary_index]:
                text = para.text.strip()
                if not text:
                    continue
                if self._normalize_title_token(text).startswith(self._normalize_title_token(self.spec.abstract.keywords_label)):
                    keyword_count = len(self._extract_keywords(text, (self.spec.abstract.keywords_label,)))
                    continue
                content_parts.append(text)

            content = "".join(content_parts).replace(" ", "")
            if len(content) < self.spec.abstract.min_chars:
                report.add_issue(
                    FormatIssueLevel.WARNING,
                    "中文摘要",
                    f"摘要内容可能过短（{len(content)}字），建议不少于{self.spec.abstract.min_chars}字",
                    "补充摘要内容",
                    f"{len(content)}字",
                    f"不少于{self.spec.abstract.min_chars}字",
                )
            if keyword_count and not 3 <= keyword_count <= 5:
                report.add_issue(
                    FormatIssueLevel.WARNING,
                    "中文摘要",
                    f"中文关键词数量为{keyword_count}个，模板要求3-5个",
                    "调整关键词数量",
                    f"{keyword_count}个",
                    "3-5个",
                )

        if self.spec.abstract.require_english_abstract:
            if english_index == -1:
                report.add_issue(FormatIssueLevel.WARNING, "文档结构", "未检测到英文摘要（Abstract）", "请补充英文摘要页")
            else:
                boundary_index = toc_index if toc_index > english_index else len(paragraphs)
                english_keyword_count = 0
                for para in paragraphs[english_index + 1:boundary_index]:
                    text = para.text.strip()
                    if not text:
                        continue
                    english_keyword_count = len(
                        self._extract_keywords(
                            text,
                            (self.spec.abstract.english_keywords_label, "Keywords", "Key words"),
                        )
                    )
                    if english_keyword_count:
                        break
                if english_keyword_count and not 3 <= english_keyword_count <= 5:
                    report.add_issue(
                        FormatIssueLevel.WARNING,
                        "英文摘要",
                        f"英文关键词数量为{english_keyword_count}个，模板要求3-5个",
                        "调整英文关键词数量",
                        f"{english_keyword_count}个",
                        "3-5个",
                    )

    def _check_references(self, doc: Document, report: FormatReport) -> None:
        ref_found = False
        ref_section_started = False
        ref_items: list[tuple[int, str]] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            text_ns = self._normalize_title_token(text)
            if text_ns in {"参考文献", "references"} and len(text_ns) < 20:
                ref_found = True
                ref_section_started = True
                for run in para.runs:
                    if run.text.strip() and run.font.name not in {None, self.spec.fonts.heading_font}:
                        report.add_issue(
                            FormatIssueLevel.WARNING, "参考文献标题",
                            f"'参考文献'标题字体应为{self.spec.fonts.heading_font}",
                            f"修改字体为{self.spec.fonts.heading_font}",
                            run.font.name or "未设置", self.spec.fonts.heading_font,
                        )
                continue
            if ref_section_started and text and len(text) > 10:
                ref_items.append((0, text))
            if ref_section_started and text_ns in {self._normalize_title_token(item) for item in self.spec.references.section_end_keywords}:
                break

        if not ref_found:
            report.add_issue(FormatIssueLevel.WARNING, "文档结构", "未检测到参考文献章节", "请添加参考文献")
            return

        self._analyze_reference_format(ref_items, report)

    def _analyze_reference_format(self, ref_items: list[tuple[int, str]], report: FormatReport) -> None:
        if not ref_items:
            report.add_issue(FormatIssueLevel.WARNING, "参考文献", "参考文献章节无具体条目", "请添加参考文献条目")
            return

        journal = book = conf = other = 0
        foreign = recent = 0
        issues = []
        current_year = datetime.now().year
        recent_start_year = current_year - self.spec.references.recent_year_span + 1
        for idx, (_, text) in enumerate(ref_items[:20]):
            if re.search(r"[\u4e00-\u9fa5]+.*\d{4}.*\d+.*\d+.*\d+", text) and ("卷" in text or "期" in text):
                journal += 1
                if not re.search(r"\d{4}.*[\(（]\d+[\)）]", text):
                    issues.append(f"第{idx+1}条期刊文献可能缺少期号格式")
            elif re.search(r"(出版社|出版|Press|Publishing)", text, re.IGNORECASE) and re.search(r"\d{4}", text):
                book += 1
            elif re.search(r"(会议|论文集|Proceedings|Conference|Symposium)", text, re.IGNORECASE):
                conf += 1
            else:
                other += 1

        for _, text in ref_items:
            if not re.search(r"[\u4e00-\u9fa5]", text):
                foreign += 1
            years = [int(year) for year in re.findall(r"\b(19\d{2}|20\d{2})\b", text)]
            if years and max(years) >= recent_start_year:
                recent += 1

        total = len(ref_items)
        report.add_issue(
            FormatIssueLevel.INFO, "参考文献统计",
            (
                f"检测到参考文献共{total}条：期刊约{journal}条，图书约{book}条，会议约{conf}条，其他约{other}条；"
                f"外文文献约{foreign}条，近{self.spec.references.recent_year_span}年（{recent_start_year}-{current_year}年）文献约{recent}条"
            ),
            "",
        )

        for iss in issues[:5]:
            report.add_issue(FormatIssueLevel.WARNING, "参考文献格式", iss, f"请参考{self.spec.references.standard}标准格式")

        if total < self.spec.references.min_count:
            report.add_issue(
                FormatIssueLevel.WARNING, "参考文献数量",
                f"参考文献数量较少（{total}条），模板要求不少于{self.spec.references.min_count}篇",
                "建议补充相关文献", f"{total}条", f"不少于{self.spec.references.min_count}条",
            )
        elif total > self.spec.references.max_typical_count:
            report.add_issue(FormatIssueLevel.INFO, "参考文献数量", f"参考文献数量较多（{total}条），请确保质量", "")

        if self.spec.references.min_foreign_count and foreign < self.spec.references.min_foreign_count:
            report.add_issue(
                FormatIssueLevel.WARNING,
                "参考文献数量",
                f"外文文献数量较少（{foreign}条），模板要求不少于{self.spec.references.min_foreign_count}篇",
                "请补充外文参考文献",
                f"{foreign}条",
                f"不少于{self.spec.references.min_foreign_count}条",
            )

        if self.spec.references.min_recent_count and recent < self.spec.references.min_recent_count:
            report.add_issue(
                FormatIssueLevel.WARNING,
                "参考文献数量",
                (
                    f"近{self.spec.references.recent_year_span}年（{recent_start_year}-{current_year}年）文献较少"
                    f"（{recent}条），模板要求不少于{self.spec.references.min_recent_count}篇"
                ),
                "请补充近年参考文献",
                f"{recent}条",
                f"不少于{self.spec.references.min_recent_count}条",
            )

    def _check_signature(self, doc: Document, report: FormatReport) -> None:
        found = False
        front_text = self._collect_front_text(doc, max_paragraphs=60)
        if all(keyword in front_text for keyword in ("专业", "学生姓名", "指导教师")):
            found = True
        else:
            for idx, para in enumerate(doc.paragraphs[:10]):
                if self._is_signature_paragraph(para.text, idx):
                    found = True
                    break
        if not found:
            report.add_issue(FormatIssueLevel.INFO, "文档结构", "未检测到署名（年级专业、学生姓名、指导教师）", "毕业论文需要包含署名")

    def _check_cover_page(self, doc: Document, report: FormatReport) -> None:
        """检查封面页是否包含 spec.cover 中定义的必要字段标签。"""
        cv = self.spec.cover
        missing = []
        first_text = self._collect_front_text(doc, max_paragraphs=40)
        for label, _ in cv.fields:
            if label not in first_text:
                missing.append(label)
        if missing:
            report.add_issue(
                FormatIssueLevel.WARNING,
                "封面页",
                f"封面页可能缺少以下字段：{', '.join(missing)}",
                "请检查封面页是否包含所有必要信息",
            )

    def _check_declaration_page(self, doc: Document, report: FormatReport) -> None:
        cv = self.spec.cover
        if not cv.declaration_title:
            return

        first_text = self._collect_front_text(doc, max_paragraphs=80)
        if cv.declaration_title not in first_text:
            report.add_issue(
                FormatIssueLevel.WARNING,
                "文档结构",
                f"未检测到“{cv.declaration_title}”页面",
                "请在封面后补充原创性保证书页面",
            )
            return

        missing_fields = [field for field in cv.declaration_fields if field not in first_text]
        if missing_fields:
            report.add_issue(
                FormatIssueLevel.WARNING,
                "原创性保证书",
                f"原创性保证书可能缺少以下字段：{', '.join(missing_fields)}",
                "请补充专业、班级、签名等信息栏",
            )

    def _check_required_sections(self, doc: Document, report: FormatReport) -> None:
        if not self.spec.required_sections:
            return

        normalized_text = self._normalize_title_token(self._collect_front_text(doc, max_paragraphs=max(len(doc.paragraphs), 1)))
        missing_sections = [
            section for section in self.spec.required_sections
            if self._normalize_title_token(section) not in normalized_text
        ]
        if missing_sections:
            report.add_issue(
                FormatIssueLevel.WARNING,
                "文档结构",
                f"模板要求的章节可能缺失：{', '.join(missing_sections)}",
                "请按模板补齐前置页和主要章节",
            )

    # ── 参考文献格式化 ────────────────────────────────────

    def _format_references(self, doc: Document, report: FormatReport) -> None:
        ref_section_started = False
        ref_items: list[tuple[object, str]] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            text_ns = text.replace(" ", "")
            if ("参考文献" in text_ns or "References" in text) and len(text) < 20:
                ref_section_started = True
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in para.runs:
                    run.font.name = self.spec.references.title_font
                    run.font.size = Pt(self.spec.references.title_size)
                    run.font.bold = True
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), self.spec.references.title_font)
                continue
            if ref_section_started and text and len(text) > 10:
                clean = re.sub(r"^\[\d+\]\s*", "", text)
                clean = re.sub(r"^\d+[.．]\s*", "", clean)
                ref_items.append((para, clean))
            if ref_section_started and text and len(text) < 10:
                if any(kw in text for kw in self.spec.references.section_end_keywords):
                    break

        if ref_items:
            numbering_ready = self._create_reference_numbering(doc)
        else:
            numbering_ready = False

        for idx, (para, clean_text) in enumerate(ref_items, 1):
            para.clear()
            text_to_write = clean_text if numbering_ready else f"[{idx}] {clean_text}"
            self._add_reference_text_with_mixed_fonts(para, text_to_write)
            para.paragraph_format.line_spacing = self._resolve_body_line_spacing()
            para.paragraph_format.first_line_indent = Cm(-self.spec.references.hanging_indent)
            para.paragraph_format.left_indent = Cm(self.spec.references.hanging_indent)
            para.paragraph_format.space_after = Pt(3)
            if numbering_ready:
                self._apply_reference_numbering(para, idx)

    def _apply_reference_numbering(self, paragraph, number: int) -> None:
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        pPr = paragraph._p.get_or_add_pPr()
        num_id = str(self.spec.references.numbering_id)
        numPr = parse_xml(
            f'<w:numPr {nsdecls("w")}><w:ilvl w:val="0"/><w:numId w:val="{num_id}"/></w:numPr>'
        )
        for child in list(pPr):
            if child.tag.endswith("numPr"):
                pPr.remove(child)
        pPr.append(numPr)
        paragraph.paragraph_format.first_line_indent = None
        paragraph.paragraph_format.left_indent = None

    def _add_reference_text_with_mixed_fonts(self, paragraph, text: str) -> None:
        current_text = ""
        is_current_chinese = None
        fs = self.spec.fonts
        body_size = self.spec.font_sizes.body

        for char in text:
            is_chinese = "\u4e00" <= char <= "\u9fff"
            if is_current_chinese is None:
                is_current_chinese = is_chinese
                current_text = char
            elif is_current_chinese == is_chinese:
                current_text += char
            else:
                if current_text:
                    run = paragraph.add_run(current_text)
                    run.font.size = Pt(body_size)
                    if is_current_chinese:
                        run.font.name = fs.chinese_font
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), fs.chinese_font)
                    else:
                        run.font.name = fs.english_font
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), fs.english_font)
                is_current_chinese = is_chinese
                current_text = char

        if current_text:
            run = paragraph.add_run(current_text)
            run.font.size = Pt(body_size)
            if is_current_chinese:
                run.font.name = fs.chinese_font
                run._element.rPr.rFonts.set(qn("w:eastAsia"), fs.chinese_font)
            else:
                run.font.name = fs.english_font
                run._element.rPr.rFonts.set(qn("w:eastAsia"), fs.english_font)

    def _create_reference_numbering(self, doc: Document) -> bool:
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        try:
            numbering_part = doc.part.numbering_part
        except Exception:
            return False

        num_id = str(self.spec.references.numbering_id)
        numbering_element = numbering_part._element
        xml_text = numbering_element.xml
        if f'w:numId="{num_id}"' in xml_text:
            return True

        try:
            abstract_num = parse_xml(
                f'<w:abstractNum {nsdecls("w")} w:abstractNumId="{num_id}">'
                '<w:multiLevelType w:val="singleLevel"/>'
                '<w:lvl w:ilvl="0">'
                '<w:numFmt w:val="decimal"/>'
                '<w:lvlText w:val="[%1]"/>'
                '<w:lvlJc w:val="left"/>'
                '<w:pPr><w:ind w:left="0" w:firstLine="0"/></w:pPr>'
                '<w:rPr><w:rFonts w:ascii="Times New Roman" w:eastAsia="Times New Roman"/></w:rPr>'
                '</w:lvl>'
                '</w:abstractNum>'
            )
            num_elem = parse_xml(
                f'<w:num {nsdecls("w")} w:numId="{num_id}">'
                f'<w:abstractNumId w:val="{num_id}"/>'
                '<w:lvlOverride w:ilvl="0"><w:startOverride w:val="1"/></w:lvlOverride>'
                '</w:num>'
            )
            numbering_element.append(abstract_num)
            numbering_element.append(num_elem)
            return True
        except Exception:
            return False

    # ── 报告参考信息 ──────────────────────────────────────

    def _get_format_reference(self) -> str:
        s = self.spec
        p = s.page_layout
        abstract_font = s.fonts.abstract_font
        h1_align = s.headings[1].alignment
        h1_pos = "居中" if h1_align == "center" else "靠左顶格"
        if s.spacing.body_line_spacing_rule == "exact":
            body_spacing_desc = f"固定值 {s.spacing.body_line_spacing}pt"
        else:
            body_spacing_desc = f"{s.spacing.body_line_spacing}倍行距"
        return f"""
【学校模板】{s.school_name}

【页面设置】
- 纸张: A4
- 上边距: {p.top_margin}cm, 下边距: {p.bottom_margin}cm
- 左边距: {p.left_margin}cm, 右边距: {p.right_margin}cm
- 装订线: {p.gutter}cm
- 页眉: {p.header_distance}cm, 页脚: {p.footer_distance}cm

【字体字号】
- 论文题目: {s.headings[0].font} {s.font_sizes.title}pt, 居中
- 一级标题: {s.headings[1].font} {s.font_sizes.heading_1}pt, {h1_pos}
- 二级标题: {s.headings[2].font} {s.font_sizes.heading_2}pt, 靠左顶格
- 三级标题: {s.headings[3].font} {s.font_sizes.heading_3}pt, 靠左顶格
- 正文: {s.fonts.chinese_font} {s.font_sizes.body}pt, {body_spacing_desc}
- 英文/数字: {s.fonts.english_font}
- 摘要内容: {abstract_font} {s.font_sizes.body}pt
- 署名: {s.signature.font} {s.signature.size}pt, 居中

【摘要要求】
- 中文摘要约300字
- "摘要"二字用{s.abstract.label_font} {s.abstract.font_size}pt, 居中
- 内容用{abstract_font} {s.font_sizes.body}pt

【参考文献】
- 格式标准: {s.references.standard}
- 数量要求: 不少于{s.references.min_count}篇
- 外文文献: 不少于{s.references.min_foreign_count}篇
- 近{ s.references.recent_year_span }年文献: 不少于{s.references.min_recent_count}篇
- 期刊格式: 作者.文章题目名[J].期刊名,年份,卷号(期数):页码
- 图书格式: 作者.书名[M].出版地:出版单位,年份:页码
- 会议格式: 作者.文章题目名[C].会议名(论文集),年份:页码
"""
