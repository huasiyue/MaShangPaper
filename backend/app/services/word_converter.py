#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将 Markdown 格式的论文初稿转换为 Word 文档（支持数学公式）
使用标准 Word Heading 样式，方便后续在 Word 中直接修改样式
公式使用 OMML 格式，可在 Word 中直接编辑
"""

import os
import re
import sys
import argparse
from copy import deepcopy
from pathlib import Path
from urllib.parse import urlparse
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml
from app.services.math_utils import latex_to_omml
from app.services.formatters.spec import ThesisFormatSpec
from app.services.formatters.schools import get_spec as _get_spec_from_registry


MATH_FONT = 'Times New Roman'

# 对齐方式映射（与 engine.py 保持一致）
_ALIGNMENT_MAP = {
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def _resolve_spec(school_id: str, thesis_type: str) -> ThesisFormatSpec:
    """根据 school_id 从注册表加载格式规格。新增学校只需注册 spec 工厂。"""
    return _get_spec_from_registry(school_id, thesis_type)


def _normalize_title_token(text: str) -> str:
    return re.sub(r'\s+', '', text).strip().lower()


def _normalize_field_label(text: str) -> str:
    cleaned = text.replace('*', '').replace('_', '').strip()
    cleaned = re.sub(r'[：:]', '', cleaned)
    cleaned = re.sub(r'\s+', '', cleaned)
    return cleaned.lower()


_COVER_FIELD_ALIASES = {
    '题目': ('题目', '论文题目', '标题', 'title'),
    '教学机构': ('教学机构', '学院', '所在学院'),
    '专业': ('专业',),
    '年级、班级': ('年级、班级', '班级', '年级班级'),
    '学号': ('学号',),
    '学生姓名': ('学生姓名', '姓名'),
    '指导教师': ('指导教师', '导师'),
    '企业导师': ('企业导师',),
    '完成日期': ('完成日期', '日期'),
}

def _get_display_title(spec: ThesisFormatSpec, text: str) -> str:
    normalized = _normalize_title_token(text)
    for source, display in spec.special_title_display_map:
        if _normalize_title_token(source) == normalized:
            return display
    return text.strip()


def _has_heading(lines: list[str], *candidates: str) -> bool:
    normalized_candidates = {_normalize_title_token(candidate) for candidate in candidates if candidate}
    for line in lines:
        stripped = line.strip()
        match = re.match(r'^(#{1,6})\s+(.*)', stripped)
        if not match:
            continue
        title = clean_heading(match.group(2).strip())
        if _normalize_title_token(title) in normalized_candidates:
            return True
    return False


def _extract_cover_metadata(lines: list[str]) -> dict[str, str]:
    metadata: dict[str, str] = {}
    in_cover = False

    for line in lines:
        stripped = line.strip()
        heading_match = re.match(r'^(#{1,6})\s+(.*)', stripped)
        if heading_match:
            title = clean_heading(heading_match.group(2).strip())
            normalized = _normalize_title_token(title)
            if normalized in {'封面', 'cover'}:
                in_cover = True
                continue
            if in_cover:
                break

        if not in_cover or not stripped:
            continue

        candidate = stripped.replace('**', '').replace('__', '').strip()
        parts = re.split(r'[：:]', candidate, maxsplit=1)
        if len(parts) != 2:
            continue
        key = _normalize_field_label(parts[0])
        value = parts[1].strip()
        if value:
            metadata[key] = value

    return metadata


def _lookup_cover_metadata(metadata: dict[str, str], *aliases: str) -> str | None:
    normalized_aliases = {_normalize_field_label(alias) for alias in aliases if alias}
    for alias in normalized_aliases:
        value = metadata.get(alias)
        if value:
            return value
    return None


def _ensure_page_break(doc):
    if any(paragraph.text.strip() for paragraph in doc.paragraphs):
        doc.add_page_break()


def _get_style(doc, style_name: str):
    try:
        return doc.styles[style_name]
    except KeyError:
        return None


def _resolve_body_line_spacing(spec: ThesisFormatSpec):
    if spec.spacing.body_line_spacing_rule == "exact":
        return Pt(spec.spacing.body_line_spacing)
    return spec.spacing.body_line_spacing


def _resolve_body_first_line_indent(spec: ThesisFormatSpec):
    return Pt(spec.font_sizes.body * 2)


def _resolve_front_title_size(spec: ThesisFormatSpec):
    if spec.school_id == "sdfmu":
        return Pt(spec.font_sizes.title)
    return Pt(max(spec.font_sizes.abstract_label, spec.font_sizes.heading_1))


def _apply_page_layout_to_sections(doc, spec: ThesisFormatSpec):
    pg = spec.page_layout
    for section in doc.sections:
        section.top_margin = Cm(pg.top_margin)
        section.bottom_margin = Cm(pg.bottom_margin)
        section.left_margin = Cm(pg.left_margin)
        section.right_margin = Cm(pg.right_margin)
        section.gutter = Cm(pg.gutter)
        section.header_distance = Cm(pg.header_distance)
        section.footer_distance = Cm(pg.footer_distance)


def _alignment_to_anchor(value) -> str:
    if value == WD_ALIGN_PARAGRAPH.LEFT:
        return "left"
    if value == WD_ALIGN_PARAGRAPH.RIGHT:
        return "right"
    return "center"


def _set_cell_borders(cell, top: str = "single", bottom: str = "single", left: str = "single", right: str = "single",
                      top_size: str = "6", bottom_size: str = "6", left_size: str = "6", right_size: str = "6",
                      color: str = "BFBFBF"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge, val, size in (
        ("top", top, top_size),
        ("bottom", bottom, bottom_size),
        ("left", left, left_size),
        ("right", right, right_size),
    ):
        edge_el = OxmlElement(f"w:{edge}")
        edge_el.set(qn("w:val"), val)
        edge_el.set(qn("w:sz"), size)
        edge_el.set(qn("w:space"), "0")
        edge_el.set(qn("w:color"), color)
        borders.append(edge_el)
    for child in list(tc_pr):
        if child.tag.endswith("tcBorders"):
            tc_pr.remove(child)
    tc_pr.append(borders)


def _set_table_font(paragraph, font_name: str, font_size: float, bold: bool = False):
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(4)
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.bold = bold
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def _set_cell_margins(cell, top: int = 0, bottom: int = 0, left: int = 0, right: int = 0):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for edge, size in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        edge_el = OxmlElement(f"w:{edge}")
        edge_el.set(qn("w:w"), str(size))
        edge_el.set(qn("w:type"), "dxa")
        tc_mar.append(edge_el)
    for child in list(tc_pr):
        if child.tag.endswith("tcMar"):
            tc_pr.remove(child)
    tc_pr.append(tc_mar)


def _try_add_cover_logo(doc, spec: ThesisFormatSpec):
    cv = spec.cover
    if not cv.logo_path or cv.logo_width <= 0:
        return

    logo_path = Path(cv.logo_path)
    if not logo_path.exists():
        return

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(12)
    try:
        _add_top_bottom_wrapped_picture(
            run,
            str(logo_path),
            width=Cm(cv.logo_width),
            height=Cm(cv.logo_height) if cv.logo_height > 0 else None,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            image_name=logo_path.name,
        )
    except Exception:
        return


def _convert_inline_shape_to_top_bottom_wrap(shape, alignment: str = "center"):
    inline = shape._inline
    inline_children = list(inline)
    anchor = OxmlElement("wp:anchor")
    anchor.set("simplePos", "0")
    anchor.set("relativeHeight", "0")
    anchor.set("behindDoc", "0")
    anchor.set("locked", "0")
    anchor.set("layoutInCell", "1")
    anchor.set("allowOverlap", "0")
    anchor.set("distT", "0")
    anchor.set("distB", "0")
    anchor.set("distL", "0")
    anchor.set("distR", "0")

    simple_pos = OxmlElement("wp:simplePos")
    simple_pos.set("x", "0")
    simple_pos.set("y", "0")
    anchor.append(simple_pos)

    position_h = OxmlElement("wp:positionH")
    position_h.set("relativeFrom", "margin")
    align_h = OxmlElement("wp:align")
    align_h.text = alignment
    position_h.append(align_h)
    anchor.append(position_h)

    position_v = OxmlElement("wp:positionV")
    position_v.set("relativeFrom", "paragraph")
    pos_offset_v = OxmlElement("wp:posOffset")
    pos_offset_v.text = "0"
    position_v.append(pos_offset_v)
    anchor.append(position_v)

    anchor.append(deepcopy(inline.extent))

    effect_extent = OxmlElement("wp:effectExtent")
    effect_extent.set("l", "0")
    effect_extent.set("t", "0")
    effect_extent.set("r", "0")
    effect_extent.set("b", "0")
    anchor.append(effect_extent)

    anchor.append(OxmlElement("wp:wrapTopAndBottom"))
    anchor.append(deepcopy(inline.docPr))
    if len(inline_children) >= 3:
        anchor.append(deepcopy(inline_children[2]))
    anchor.append(deepcopy(inline.graphic))

    drawing = inline.getparent()
    drawing.remove(inline)
    drawing.append(anchor)


def _add_top_bottom_wrapped_picture(run, image_source, width, alignment, image_name: str = "Picture", height=None):
    shape = run.add_picture(image_source, width=width, height=height)
    # 同一文档中可能会插入多张图片，这里复用 python-docx 分配的 docPr 信息，再改为浮动锚点。
    _convert_inline_shape_to_top_bottom_wrap(shape, alignment=_alignment_to_anchor(alignment))
    return shape


def _add_block_equation(doc, latex_str, eq_label: str = ""):
    """向 Word 文档中添加一个居中的块级 OMML 公式段落。

    使用三列不可见表格实现：左占位 | 居中公式 | 右编号。
    """
    if not eq_label:
        # 无编号：直接居中段落
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.line_spacing = 1.5
        try:
            omml_elem = latex_to_omml(latex_str, display='block')
            para._element.append(omml_elem)
        except Exception:
            run = para.add_run(latex_str)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        return

    # 有编号：三列不可见表格，公式居中，编号靠右
    tbl = doc.add_table(rows=1, cols=3)
    tbl.autofit = True

    # 移除表格所有边框和样式
    tbl_style = tbl.style
    try:
        tbl.style = 'Table Grid'
    except Exception:
        pass

    # 设置表格宽度为页面可用宽度
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    tbl_element = tbl._tbl
    tblPr = tbl_element.tblPr

    # 无边框
    tblBorders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    )
    # 移除旧边框
    for child in list(tblPr):
        if child.tag.endswith('tblBorders'):
            tblPr.remove(child)
    tblPr.append(tblBorders)

    # 中间列放公式（居中）
    cell_eq = tbl.cell(0, 1)
    cell_eq.width = Cm(10)
    para_eq = cell_eq.paragraphs[0]
    para_eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_eq.paragraph_format.line_spacing = 1.5
    try:
        omml_elem = latex_to_omml(latex_str, display='block')
        para_eq._element.append(omml_elem)
    except Exception:
        run = para_eq.add_run(latex_str)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    # 右列放编号
    cell_num = tbl.cell(0, 2)
    cell_num.width = Cm(2)
    para_num = cell_num.paragraphs[0]
    para_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_num = para_num.add_run(eq_label)
    run_num.font.name = 'Times New Roman'
    run_num.font.size = Pt(12)


def _add_inline_equation(para, latex_str):
    """向已有段落中插入一个行内 OMML 公式。"""
    try:
        omml_elem = latex_to_omml(latex_str, display='inline')
        para._element.append(omml_elem)
    except Exception:
        run = para.add_run(f'${latex_str}$')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)


def set_math_font(doc):
    """设置文档的默认公式字体为 Times New Roman。"""
    from lxml import etree
    M = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
    # 在 settings.xml 中添加/修改 <m:mathPr><m:mathFont>
    settings = doc.settings.element
    mathPr = settings.find(qn('m:mathPr'))
    if mathPr is None:
        mathPr = OxmlElement('m:mathPr')
        settings.append(mathPr)
    mathFont = mathPr.find(qn('m:mathFont'))
    if mathFont is None:
        mathFont = OxmlElement('m:mathFont')
        mathPr.insert(0, mathFont)
    mathFont.set(qn('m:val'), MATH_FONT)


def add_page_break(paragraph):
    """在段落前添加分页符"""
    run = paragraph.add_run()
    run._r.append(OxmlElement('w:br'))
    run._r[-1].set(qn('w:type'), 'page')


def _insert_toc_field(doc):
    """在当前位置插入 Word 目录域（TOC field）。

    生成一个包含 TOC 域代码的段落，Word 打开时会自动更新目录。
    """
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # 创建域代码结构: <w:fldSimple w:instr=" TOC \\o &quot;1-3&quot; \\h \\z \\u ">
    fldSimple = OxmlElement('w:fldSimple')
    fldSimple.set(qn('w:instr'), ' TOC \\o "1-3" \\h \\z \\u ')

    # 域内占位文字
    fldContent = OxmlElement('w:fldContent')
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), '宋体')
    rFonts.set(qn('w:eastAsia'), '宋体')
    rFonts.set(qn('w:hAnsi'), '宋体')
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '24')  # 12pt
    rPr.append(sz)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = '（请在 Word 中右键点击此处，选择"更新域"以生成目录）'
    r.append(t)
    fldContent.append(r)
    fldSimple.append(fldContent)
    para._element.append(fldSimple)
    return para


def _add_cover_blank_line(doc, line_spacing: float | None = None):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.first_line_indent = Pt(0)
    para.paragraph_format.left_indent = Pt(0)
    para.paragraph_format.right_indent = Pt(0)
    if line_spacing is not None:
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        para.paragraph_format.line_spacing = line_spacing
    run = para.add_run("")
    run.font.name = "仿宋_GB2312"
    run.font.size = Pt(16)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")
    return para


def _add_cover_gap_line(doc):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run("")
    run.font.name = "仿宋_GB2312"
    run.font.size = Pt(16)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")
    return para


def _add_cover_date_gap_line(doc):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run("")
    run.font.name = "仿宋_GB2312"
    run.font.size = Pt(16)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")
    return para


def _generate_cover_page(doc, spec: ThesisFormatSpec, metadata: dict[str, str] | None = None):
    """生成论文封面页模板，尽量贴近学校模板版式。"""
    from docx.enum.table import WD_TABLE_ALIGNMENT
    cv = spec.cover
    metadata = metadata or {}

    _add_cover_blank_line(doc, line_spacing=1.25)

    # 学校名称
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(cv.university_name)
    run.font.name = cv.university_font
    run.font.size = Pt(cv.university_size)
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), cv.university_font)

    # 论文类型标签（如"本科毕业论文（设计）"）
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(cv.thesis_type_label)
    run.font.name = cv.thesis_type_font
    run.font.size = Pt(cv.thesis_type_size)
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), cv.thesis_type_font)

    _add_cover_blank_line(doc, line_spacing=1.25)
    _try_add_cover_logo(doc, spec)

    # 论文题目区域
    title_table = doc.add_table(rows=1, cols=2)
    title_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    title_table.autofit = False
    title_table.columns[0].width = Cm(cv.title_table_left_width or 3.44)
    title_table.columns[1].width = Cm(cv.title_table_right_width or 13.54)
    title_left = title_table.cell(0, 0)
    title_right = title_table.cell(0, 1)
    row = title_table.rows[0]
    if cv.title_table_row_height > 0:
        row.height = Cm(cv.title_table_row_height)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    title_left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    title_right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    for cell in (title_left, title_right):
        _set_cell_borders(cell, top="dashed", bottom="dashed", left="dashed", right="dashed", color="BFBFBF")
        _set_cell_margins(cell, top=0, bottom=0, left=0, right=0)

    left_para = title_left.paragraphs[0]
    left_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    left_para.paragraph_format.space_before = Pt(0)
    left_para.paragraph_format.space_after = Pt(0)
    left_para.paragraph_format.line_spacing = Pt(23)
    left_run = left_para.add_run("题目：")
    left_run.font.name = cv.title_font
    left_run.font.size = Pt(cv.title_size)
    left_run.font.bold = False
    left_run.font.underline = False
    left_run._element.rPr.rFonts.set(qn('w:eastAsia'), cv.title_font)

    p = title_right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(23)
    p.paragraph_format.left_indent = Pt(0)
    p.paragraph_format.right_indent = Pt(0)
    p.paragraph_format.first_line_indent = Pt(0)
    title_text = _lookup_cover_metadata(metadata, *_COVER_FIELD_ALIASES['题目']) or '（论文题目）'
    run = p.add_run(title_text)
    run.font.name = cv.title_font
    run.font.size = Pt(cv.title_size)
    run.font.bold = False
    run.font.underline = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), cv.title_font)

    _add_cover_gap_line(doc)

    # 字段表格
    table = doc.add_table(rows=len(cv.fields), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(cv.info_table_left_width or 3.2)
    table.columns[1].width = Cm(cv.info_table_right_width or 9.4)

    for i, (label, placeholder) in enumerate(cv.fields):
        row = table.rows[i]
        if cv.info_table_row_height > 0:
            row.height = Cm(cv.info_table_row_height)
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        left_cell = row.cells[0]
        right_cell = row.cells[1]
        left_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        right_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        field_value = _lookup_cover_metadata(metadata, *_COVER_FIELD_ALIASES.get(label, (label,))) or placeholder
        _set_cell_borders(left_cell, top="dashed", bottom="dashed", left="dashed", right="dashed", color="BFBFBF")
        _set_cell_borders(
            right_cell,
            top="single",
            bottom="single",
            left="none",
            right="single",
            color="000000",
        )

        # 左列：标签
        p_left = left_cell.paragraphs[0]
        p_left.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_label = p_left.add_run(label)
        _set_table_font(p_left, "仿宋_GB2312", 16)

        # 右列：内容
        p_right = right_cell.paragraphs[0]
        p_right.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_underline = p_right.add_run(field_value)
        _set_table_font(p_right, "仿宋_GB2312", 16)

    _add_cover_date_gap_line(doc)
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.paragraph_format.space_before = Pt(0)
    date_text = _lookup_cover_metadata(metadata, *_COVER_FIELD_ALIASES.get("完成日期", ("日期",))) or "____年__月__日"
    date_run = date_para.add_run(date_text)
    date_run.font.name = "仿宋_GB2312"
    date_run.font.size = Pt(16)
    date_run._element.rPr.rFonts.set(qn('w:eastAsia'), "仿宋_GB2312")


def _generate_declaration_page(doc, spec: ThesisFormatSpec, metadata: dict[str, str] | None = None):
    """生成原创性保证书页面。"""
    from docx.enum.table import WD_TABLE_ALIGNMENT
    cv = spec.cover
    metadata = metadata or {}
    if not cv.declaration_title:
        return

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(70)
    title.paragraph_format.space_after = Pt(22)
    run = title.add_run(cv.declaration_title)
    run.font.name = spec.fonts.heading_font
    run.font.size = Pt(spec.font_sizes.title - 2)
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), spec.fonts.heading_font)

    for line in cv.declaration_body:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.line_spacing = _resolve_body_line_spacing(spec)
        para.paragraph_format.first_line_indent = Pt(0)
        para.paragraph_format.space_after = Pt(6)
        run = para.add_run(line)
        run.font.name = spec.fonts.english_font
        run.font.size = Pt(spec.font_sizes.body + 2)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), spec.fonts.chinese_font)

    doc.add_paragraph()

    table = doc.add_table(rows=len(cv.declaration_fields), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(2.0)
    table.columns[1].width = Cm(8.2)

    for row_idx, field in enumerate(cv.declaration_fields):
        left_cell = table.cell(row_idx, 0)
        right_cell = table.cell(row_idx, 1)
        _set_cell_borders(left_cell, color="BFBFBF")
        _set_cell_borders(right_cell, left="none", color="000000")

        if field == "专业":
            field_value = _lookup_cover_metadata(metadata, *_COVER_FIELD_ALIASES.get("专业", ("专业",)))
        elif field == "班级":
            field_value = _lookup_cover_metadata(metadata, *_COVER_FIELD_ALIASES.get("年级、班级", ("班级",)))
        else:
            field_value = None

        label_para = left_cell.paragraphs[0]
        label_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        label_para.paragraph_format.space_before = Pt(4)
        label_para.paragraph_format.space_after = Pt(4)
        label_run = label_para.add_run(f"{field}：")
        label_run.font.name = spec.fonts.english_font
        label_run.font.size = Pt(spec.font_sizes.body)
        label_run._element.rPr.rFonts.set(qn('w:eastAsia'), spec.fonts.chinese_font)

        value_para = right_cell.paragraphs[0]
        value_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        value_para.paragraph_format.space_before = Pt(4)
        value_para.paragraph_format.space_after = Pt(4)
        value_run = value_para.add_run(field_value or "")
        value_run.font.name = spec.fonts.english_font
        value_run.font.size = Pt(spec.font_sizes.body)
        value_run._element.rPr.rFonts.set(qn('w:eastAsia'), spec.fonts.chinese_font)

    if cv.declaration_date_placeholder:
        doc.add_paragraph()
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para.paragraph_format.space_before = Pt(12)
        date_text = _lookup_cover_metadata(metadata, *_COVER_FIELD_ALIASES.get("完成日期", ("日期",))) or cv.declaration_date_placeholder
        date_run = date_para.add_run(date_text)
        date_run.font.name = spec.fonts.english_font
        date_run.font.size = Pt(spec.font_sizes.body)
        date_run._element.rPr.rFonts.set(qn('w:eastAsia'), spec.fonts.chinese_font)


def _add_section_break(doc):
    """添加分节符（下一页），用于罗马/阿拉伯双页码。"""
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml

    # 在最后一段后添加分节符
    sectPr = OxmlElement('w:sectPr')
    sectType = OxmlElement('w:type')
    sectType.set(qn('w:val'), 'nextPage')
    sectPr.append(sectType)

    # 添加到文档最后一个段落的 pPr 中，或直接加到 document body
    body = doc.element.body
    # 在最终 sectPr 之前插入
    final_sectPr = body.find(qn('w:sectPr'))
    if final_sectPr is not None:
        body.insert(list(body).index(final_sectPr), sectPr)
    else:
        body.append(sectPr)
    return sectPr


def _set_page_number_roman(section):
    """设置某节的页码格式为罗马数字（i, ii, iii）。"""
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn('w:pgNumType'))
    if pgNumType is None:
        pgNumType = OxmlElement('w:pgNumType')
        sectPr.append(pgNumType)
    pgNumType.set(qn('w:fmt'), 'lowerRoman')
    pgNumType.set(qn('w:start'), '1')


def _set_page_number_arabic(section):
    """设置某节的页码格式为阿拉伯数字，从 1 开始。"""
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn('w:pgNumType'))
    if pgNumType is None:
        pgNumType = OxmlElement('w:pgNumType')
        sectPr.append(pgNumType)
    pgNumType.set(qn('w:fmt'), 'decimal')
    pgNumType.set(qn('w:start'), '1')


def _clear_section_header_footer(section):
    header = section.header
    header.is_linked_to_previous = False
    for paragraph in header.paragraphs:
        paragraph.clear()

    footer = section.footer
    footer.is_linked_to_previous = False
    for paragraph in footer.paragraphs:
        paragraph.clear()


def _add_page_number_to_footer(section, spec: ThesisFormatSpec):
    """向页脚添加居中页码。"""
    from docx.shared import Emu
    ft = spec.fonts
    fs = spec.font_sizes

    footer = section.footer
    footer.is_linked_to_previous = False

    # 清空现有内容
    for p in footer.paragraphs:
        p.clear()

    if not footer.paragraphs:
        para = footer.add_paragraph()
    else:
        para = footer.paragraphs[0]

    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加 PAGE 域代码
    run = para.add_run()
    run.font.name = ft.english_font
    run.font.size = Pt(fs.header_footer)

    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar_begin)

    run2 = para.add_run()
    run2.font.name = ft.english_font
    run2.font.size = Pt(fs.header_footer)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run2._r.append(instrText)

    run3 = para.add_run()
    run3.font.name = ft.english_font
    run3.font.size = Pt(fs.header_footer)
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    run3._r.append(fldChar_end)


def _add_header(section, spec: ThesisFormatSpec, add_bottom_border: bool = True):
    """向页眉添加学校名称文本（右对齐，宋体 10.5pt），并添加底部下划线边框。"""
    ft = spec.fonts
    fs = spec.font_sizes

    header = section.header
    header.is_linked_to_previous = False

    if not header.paragraphs:
        para = header.add_paragraph()
    else:
        para = header.paragraphs[0]

    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    para.paragraph_format.line_spacing = 1.0
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.first_line_indent = Pt(0)
    para.paragraph_format.left_indent = Pt(0)
    para.paragraph_format.right_indent = Pt(0)
    para.clear()
    run = para.add_run(spec.header_text)
    run.font.name = ft.chinese_font
    run.font.size = Pt(fs.header_footer)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), ft.chinese_font)

    # 添加页眉底部边框
    pPr = para._p.get_or_add_pPr()
    for child in list(pPr):
        if child.tag.endswith('pBdr'):
            pPr.remove(child)
    if add_bottom_border:
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '0')
        bottom.set(qn('w:color'), '000000')
        pBdr.append(bottom)
        pPr.append(pBdr)


def parse_image_meta(alt_text: str):
    """支持 `caption|width=8cm|align=center` 语法。"""
    parts = [part.strip() for part in alt_text.split('|') if part.strip()]
    caption = parts[0] if parts else "（请在此补充图名）"
    width = Cm(12)
    alignment = WD_ALIGN_PARAGRAPH.CENTER

    for option in parts[1:]:
        if '=' not in option:
            continue

        key, value = [item.strip().lower() for item in option.split('=', 1)]
        if key in {'width', 'w'}:
            width_match = re.match(r'^(\d+(?:\.\d+)?)(cm|mm|in|pt)?$', value)
            if not width_match:
                continue

            number = float(width_match.group(1))
            unit = width_match.group(2) or 'cm'
            if unit == 'cm':
                width = Cm(number)
            elif unit == 'mm':
                width = Cm(number / 10)
            elif unit == 'in':
                width = Inches(number)
            elif unit == 'pt':
                width = Pt(number)
        elif key in {'align', 'a'}:
            if value in {'left', '左'}:
                alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif value in {'right', '右'}:
                alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                alignment = WD_ALIGN_PARAGRAPH.CENTER

    return caption, width, alignment


def set_table_borders(table, border_type='three_line'):
    """
    设置表格边框为三线表格式
    三线表：顶线、栏目线、底线
    """
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    
    # 获取表格的 tbl 元素
    tbl = table._tbl
    
    # 创建表格属性
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(r'<w:tblPr {}></w:tblPr>'.format(nsdecls('w')))
    
    # 清除所有边框
    tblBorders = parse_xml(r'''
        <w:tblBorders {}>
            <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>
            <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>
            <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>
            <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>
            <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>
            <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>
        </w:tblBorders>
    '''.format(nsdecls('w')))
    
    # 移除旧的边框设置
    for child in list(tblPr):
        if child.tag.endswith('tblBorders'):
            tblPr.remove(child)
    
    tblPr.append(tblBorders)
    
    # 如果没有 tblPr，添加它
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)
    
    # 设置顶线（第一行顶部）和栏目线（第一行底部）
    if len(table.rows) > 0:
        header_row = table.rows[0]
        for cell in header_row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = parse_xml(r'''
                <w:tcBorders {}>
                    <w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/>
                    <w:bottom w:val="single" w:sz="6" w:space="0" w:color="000000"/>
                    <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>
                    <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>
                </w:tcBorders>
            '''.format(nsdecls('w')))
            # 移除旧的单元格边框设置
            for child in list(tcPr):
                if child.tag.endswith('tcBorders'):
                    tcPr.remove(child)
            tcPr.append(tcBorders)
    
    # 设置底线（最后一行底部）
    if len(table.rows) > 1:
        bottom_row = table.rows[-1]
        for cell in bottom_row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = parse_xml(r'''
                <w:tcBorders {}>
                    <w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>
                    <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>
                    <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>
                </w:tcBorders>
            '''.format(nsdecls('w')))
            # 移除旧的单元格边框设置
            for child in list(tcPr):
                if child.tag.endswith('tcBorders'):
                    tcPr.remove(child)
            tcPr.append(tcBorders)
    
    # 中间行无边框
    for row_idx in range(1, len(table.rows) - 1):
        row = table.rows[row_idx]
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = parse_xml(r'''
                <w:tcBorders {}>
                    <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>
                    <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>
                    <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>
                    <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>
                </w:tcBorders>
            '''.format(nsdecls('w')))
            # 移除旧的单元格边框设置
            for child in list(tcPr):
                if child.tag.endswith('tcBorders'):
                    tcPr.remove(child)
            tcPr.append(tcBorders)


def format_reference_text(para, text, spec: ThesisFormatSpec):
    """
    格式化参考文献文本，确保数字和英文使用指定英文字体
    中文字符使用正文字体
    """
    import re
    ft = spec.fonts
    fs = spec.font_sizes

    # 匹配模式：序号、作者、标题、期刊、年份、卷期、页码等
    # 首先处理序号 [数字]
    parts = re.split(r'(\[\d+\])', text)

    for part in parts:
        if not part:
            continue

        # 如果是序号 [数字]
        if re.match(r'^\[\d+\]$', part):
            run = para.add_run(part)
            run.font.name = ft.english_font
            run.font.size = Pt(fs.body)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), ft.english_font)
        else:
            # 处理剩余部分，区分中英文
            current_run_text = ""
            is_current_ascii = None

            for char in part:
                is_ascii = ord(char) < 128 and char.isprintable() or char.isspace()

                if is_current_ascii is None:
                    is_current_ascii = is_ascii
                    current_run_text = char
                elif is_current_ascii == is_ascii:
                    current_run_text += char
                else:
                    if current_run_text:
                        run = para.add_run(current_run_text)
                        run.font.size = Pt(fs.body)
                        if is_current_ascii:
                            run.font.name = ft.english_font
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), ft.english_font)
                        else:
                            run.font.name = ft.chinese_font
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), ft.chinese_font)
                    is_current_ascii = is_ascii
                    current_run_text = char

            if current_run_text:
                run = para.add_run(current_run_text)
                run.font.size = Pt(fs.body)
                if is_current_ascii:
                    run.font.name = ft.english_font
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), ft.english_font)
                else:
                    run.font.name = ft.chinese_font
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), ft.chinese_font)


def setup_styles(doc, spec: ThesisFormatSpec):
    """设置文档的标准样式（从 spec 读取参数）"""
    ft = spec.fonts
    fs = spec.font_sizes
    sp = spec.spacing
    body_line_spacing = _resolve_body_line_spacing(spec)
    front_title_size = _resolve_front_title_size(spec)

    # 设置正文样式
    style = doc.styles['Normal']
    style.font.name = ft.english_font
    style.font.size = Pt(fs.body)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), ft.chinese_font)
    style.paragraph_format.line_spacing = body_line_spacing
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.first_line_indent = _resolve_body_first_line_indent(spec)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # 设置 Heading 1 样式（一级标题）
    heading1 = doc.styles['Heading 1']
    h1 = spec.headings[1]
    heading1.font.name = ft.english_font
    heading1.font.size = Pt(h1.size)
    heading1.font.bold = h1.bold
    heading1.font.color.rgb = RGBColor(0, 0, 0)
    heading1._element.rPr.rFonts.set(qn('w:eastAsia'), h1.font)
    heading1.paragraph_format.alignment = _ALIGNMENT_MAP.get(h1.alignment, WD_ALIGN_PARAGRAPH.CENTER)
    heading1.paragraph_format.space_before = Pt(sp.paragraph_spacing)
    heading1.paragraph_format.space_after = Pt(sp.paragraph_spacing)
    heading1.paragraph_format.line_spacing = sp.heading_line_spacing

    # 设置 Heading 2 样式（二级标题）
    heading2 = doc.styles['Heading 2']
    h2 = spec.headings[2]
    heading2.font.name = ft.english_font
    heading2.font.size = Pt(h2.size)
    heading2.font.bold = h2.bold
    heading2.font.color.rgb = RGBColor(0, 0, 0)
    heading2._element.rPr.rFonts.set(qn('w:eastAsia'), h2.font)
    heading2.paragraph_format.alignment = _ALIGNMENT_MAP.get(h2.alignment, WD_ALIGN_PARAGRAPH.LEFT)
    heading2.paragraph_format.space_before = Pt(sp.paragraph_spacing)
    heading2.paragraph_format.space_after = Pt(sp.paragraph_spacing)
    heading2.paragraph_format.line_spacing = sp.heading_line_spacing

    # 设置 Heading 3 样式（三级标题）
    heading3 = doc.styles['Heading 3']
    h3 = spec.headings[3]
    heading3.font.name = ft.english_font
    heading3.font.size = Pt(h3.size)
    heading3.font.bold = h3.bold
    heading3.font.color.rgb = RGBColor(0, 0, 0)
    heading3._element.rPr.rFonts.set(qn('w:eastAsia'), h3.font)
    heading3.paragraph_format.alignment = _ALIGNMENT_MAP.get(h3.alignment, WD_ALIGN_PARAGRAPH.LEFT)
    heading3.paragraph_format.space_before = Pt(sp.paragraph_spacing)
    heading3.paragraph_format.space_after = Pt(sp.paragraph_spacing)
    heading3.paragraph_format.line_spacing = sp.heading_line_spacing

    # 设置 Heading 4 样式（四级标题）
    heading4 = doc.styles['Heading 4']
    h4 = spec.headings[4]
    heading4.font.name = ft.english_font
    heading4.font.size = Pt(h4.size)
    heading4.font.bold = h4.bold
    heading4.font.color.rgb = RGBColor(0, 0, 0)
    heading4._element.rPr.rFonts.set(qn('w:eastAsia'), h4.font)
    heading4.paragraph_format.alignment = _ALIGNMENT_MAP.get(h4.alignment, WD_ALIGN_PARAGRAPH.LEFT)
    heading4.paragraph_format.space_before = Pt(sp.paragraph_spacing)
    heading4.paragraph_format.space_after = Pt(sp.paragraph_spacing)
    heading4.paragraph_format.line_spacing = sp.heading_line_spacing

    # 设置 Title 样式（摘要/Abstract 等前置标题使用，避免混入目录）
    title_builtin = _get_style(doc, 'Title')
    if title_builtin is not None:
        title_builtin.font.name = ft.english_font
        title_builtin.font.size = front_title_size
        title_builtin.font.bold = h1.bold
        title_builtin.font.color.rgb = RGBColor(0, 0, 0)
        title_builtin._element.rPr.rFonts.set(qn('w:eastAsia'), h1.font)
        title_builtin.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_builtin.paragraph_format.space_before = Pt(sp.paragraph_spacing)
        title_builtin.paragraph_format.space_after = Pt(sp.paragraph_spacing)
        title_builtin.paragraph_format.line_spacing = sp.heading_line_spacing

    # 设置 TOC Heading 样式（目录标题使用）
    toc_heading = _get_style(doc, 'TOC Heading')
    if toc_heading is not None:
        toc_heading.font.name = ft.english_font
        toc_heading.font.size = front_title_size
        toc_heading.font.bold = h1.bold
        toc_heading.font.color.rgb = RGBColor(0, 0, 0)
        toc_heading._element.rPr.rFonts.set(qn('w:eastAsia'), h1.font)
        toc_heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        toc_heading.paragraph_format.space_before = Pt(sp.paragraph_spacing)
        toc_heading.paragraph_format.space_after = Pt(sp.paragraph_spacing)
        toc_heading.paragraph_format.line_spacing = sp.heading_line_spacing

    # 设置论文标题样式（自定义样式）
    h0 = spec.headings[0]
    if '论文标题' not in doc.styles:
        title_style = doc.styles.add_style('论文标题', WD_STYLE_TYPE.PARAGRAPH)
    else:
        title_style = doc.styles['论文标题']
    title_style.font.name = ft.english_font
    title_style.font.size = Pt(h0.size)
    title_style.font.bold = h0.bold
    title_style.font.color.rgb = RGBColor(0, 0, 0)
    title_style._element.rPr.rFonts.set(qn('w:eastAsia'), h0.font)
    title_style.paragraph_format.alignment = _ALIGNMENT_MAP.get(h0.alignment, WD_ALIGN_PARAGRAPH.CENTER)
    title_style.paragraph_format.space_before = Pt(sp.paragraph_spacing)
    title_style.paragraph_format.space_after = Pt(sp.paragraph_spacing)
    title_style.paragraph_format.line_spacing = sp.title_line_spacing

    # 设置 Caption 样式（题注）
    cap = spec.caption
    if 'Caption' not in doc.styles:
        caption_style = doc.styles.add_style('Caption', WD_STYLE_TYPE.PARAGRAPH)
    else:
        caption_style = doc.styles['Caption']
    caption_style.font.name = ft.english_font
    caption_style.font.size = Pt(cap.font_size)
    caption_style.font.color.rgb = RGBColor(0, 0, 0)
    caption_style._element.rPr.rFonts.set(qn('w:eastAsia'), ft.chinese_font)
    caption_style.paragraph_format.alignment = _ALIGNMENT_MAP.get(cap.alignment, WD_ALIGN_PARAGRAPH.CENTER)
    caption_style.paragraph_format.line_spacing = cap.line_spacing
    caption_style.paragraph_format.space_before = Pt(6)
    caption_style.paragraph_format.space_after = Pt(6)


def int_to_chinese(num):
    chinese_digits = ["", "一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]
    if num == 0: return "零"
    if num <= 10: return chinese_digits[num]
    elif num < 20: return "十" + chinese_digits[num - 10]
    elif num % 10 == 0: return chinese_digits[num // 10] + "十"
    else: return chinese_digits[num // 10] + "十" + chinese_digits[num % 10]

def clean_heading(text):
    text = re.sub(r'^#{1,4}\s*', '', text)
    text = re.sub(r'^第[一二三四五六七八九十百]+章\s*', '', text)
    text = re.sub(r'^第\d+章\s*', '', text)
    text = re.sub(r'^第[一二三四五六七八九十百]+节\s*', '', text)
    text = re.sub(r'^[一二三四五六七八九十]+\s*[、\.]?\s*', '', text)
    text = re.sub(r'^（[一二三四五六七八九十]+）\s*', '', text)
    text = re.sub(r'^[\d.]+、?\s*', '', text)
    text = re.sub(r'^（\d+）\s*', '', text)
    return text.strip()

def convert_markdown_to_word(markdown_path: str, output_path: str,
                             school_id: str = "sdfmu", thesis_type: str = "thesis"):
    """将 Markdown 转换为 Word 文档，使用标准 Heading 样式"""

    # 加载学校格式规格
    spec = _resolve_spec(school_id, thesis_type)
    markdown_file = Path(markdown_path).resolve()
    markdown_base_dir = markdown_file.parent

    # 读取 Markdown 文件
    with open(markdown_file, 'r', encoding='utf-8') as f:
        content = f.read()

    # 创建 Word 文档
    doc = Document()

    # 设置页面（从 spec 读取）
    _apply_page_layout_to_sections(doc, spec)

    # 设置标准样式
    setup_styles(doc, spec)
    set_math_font(doc)

    # 解析 Markdown 内容
    lines = content.split('\n')
    cover_metadata = _extract_cover_metadata(lines)
    has_declaration_heading = _has_heading(lines, spec.cover.declaration_title)
    in_code_block = False
    in_math_block = False
    math_buffer = []
    code_buffer = []
    in_abstract = False
    in_references = False
    in_front_matter = True   # 前言部分（封面、摘要、目录）用罗马数字页码
    first_heading1 = True  # 标记是否是第一个一级标题（论文题目）
    chap_num = 0
    sec_num = 0
    subsec_num = 0
    subsubsec_num = 0
    fig_num = 0
    tab_num = 0
    eq_num = 0  # 公式编号（按章节重置）
    just_parsed_table_caption = False
    skip_auto_block_content = False
    suppress_next_page_break = False
    cover_generated = False
    declaration_generated = False
    body_section_started = False
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        
        # 跳过空行
        if not line.strip():
            i += 1
            continue
        
        stripped = line.strip()

        if skip_auto_block_content:
            if stripped.startswith('#'):
                skip_auto_block_content = False
            else:
                i += 1
                continue
        
        # 代码块处理
        if stripped.startswith('```'):
            if not in_code_block:
                # 代码块开始
                in_code_block = True
                code_buffer = []
            else:
                # 代码块结束，输出已收集的代码内容
                in_code_block = False
                if code_buffer:
                    para = doc.add_paragraph()
                    para.paragraph_format.left_indent = Cm(1)
                    run = para.add_run('\n'.join(code_buffer))
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                    code_buffer = []
            i += 1
            continue

        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        # ── 数学公式块处理 ($$...$$) ──
        # 单行公式：$$ E = mc^2 $$
        if stripped.startswith('$$') and stripped.endswith('$$') and len(stripped) > 4 and not in_math_block:
            latex = stripped[2:-2].strip()
            eq_num += 1
            eq_label = f"({chap_num}-{eq_num})" if chap_num > 0 else f"({eq_num})"
            _add_block_equation(doc, latex, eq_label)
            i += 1
            continue

        # 多行公式块开始/结束
        if stripped == '$$':
            if not in_math_block:
                in_math_block = True
                math_buffer = []
            else:
                in_math_block = False
                if math_buffer:
                    latex = '\n'.join(math_buffer)
                    eq_num += 1
                    eq_label = f"({chap_num}-{eq_num})" if chap_num > 0 else f"({eq_num})"
                    _add_block_equation(doc, latex, eq_label)
                    math_buffer = []
            i += 1
            continue

        if in_math_block:
            math_buffer.append(stripped)
            i += 1
            continue

        # -- 图片插图题注 --
        if stripped.startswith('!['):
            img_match = re.match(r'^!\[(.*?)\]\((.*?)\)$', stripped)
            if img_match:
                caption, image_width, image_alignment = parse_image_meta(img_match.group(1).strip())
                image_target = img_match.group(2).strip()
                fig_num += 1
                caption_text = f"图{chap_num}-{fig_num} {caption}" if chap_num > 0 else f"图{fig_num} {caption}"
                
                para_img = doc.add_paragraph()
                para_img.alignment = image_alignment
                run_img = para_img.add_run()

                image_path = Path(image_target)
                if not image_path.is_absolute():
                    image_path = (markdown_base_dir / image_path).resolve()
                if image_path.exists():
                    try:
                        _add_top_bottom_wrapped_picture(
                            run_img,
                            str(image_path),
                            width=image_width,
                            alignment=image_alignment,
                            image_name=image_path.name,
                        )
                    except Exception:
                        # python-docx 可能不支持某些 JPEG，用 PIL 转为 PNG 后插入
                        try:
                            from PIL import Image
                            import io
                            img = Image.open(str(image_path))
                            buf = io.BytesIO()
                            img.save(buf, format='PNG')
                            buf.seek(0)
                            _add_top_bottom_wrapped_picture(
                                run_img,
                                buf,
                                width=image_width,
                                alignment=image_alignment,
                                image_name=image_path.stem + '.png',
                            )
                        except Exception:
                            run_img = para_img.add_run(f"[插图格式不支持: {image_target}]")
                            run_img.font.color.rgb = RGBColor(128, 128, 128)
                else:
                    run_img = para_img.add_run(f"[插图: {image_target}]")
                    run_img.font.color.rgb = RGBColor(128, 128, 128)
                
                # 图题注 (图片下方)
                para = doc.add_paragraph(style='Caption')
                para.alignment = image_alignment
                para.add_run(caption_text)
                
                just_parsed_table_caption = True
                i += 1
                continue

        # -- 图/表 显示题注 --
        cap_match = re.match(r'^#*\s*(\*\*?)?(图|表)\s*(?:\d+[-\.]\d+|\d+)?[:：\s]+\s*(.+?)(\*\*?)?$', stripped)
        if cap_match and '格式' not in stripped:
            ctype = cap_match.group(2)
            caption = cap_match.group(3).strip()
            if ctype == '图':
                fig_num += 1
                caption_text = f"图{chap_num}-{fig_num} {caption}" if chap_num > 0 else f"图{fig_num} {caption}"
            else:
                tab_num += 1
                caption_text = f"表{chap_num}-{tab_num} {caption}" if chap_num > 0 else f"表{tab_num} {caption}"
            
            para = doc.add_paragraph(caption_text, style='Caption')
            just_parsed_table_caption = (ctype == '表')
            i += 1
            continue

        # -- 各类结构标题 --
        if stripped.startswith('#'):
            h_match = re.match(r'^(#{1,4})\s+(.*)', stripped)
            if h_match:
                level_marks = h_match.group(1)
                raw_title = h_match.group(2).strip()
                title_clean = clean_heading(raw_title)
                title_clean_lower = title_clean.lower()

                # ── 封面页 ──
                if '封面' in title_clean or 'cover' in title_clean_lower:
                    _generate_cover_page(doc, spec, cover_metadata)
                    cover_generated = True
                    _add_section_break(doc)
                    suppress_next_page_break = True
                    if spec.cover.declaration_title and not has_declaration_heading:
                        _generate_declaration_page(doc, spec, cover_metadata)
                        declaration_generated = True
                        _add_section_break(doc)
                        suppress_next_page_break = True
                    skip_auto_block_content = True
                    first_heading1 = False
                    i += 1
                    continue

                if spec.cover.declaration_title and _normalize_title_token(title_clean) == _normalize_title_token(spec.cover.declaration_title):
                    if not declaration_generated:
                        _generate_declaration_page(doc, spec, cover_metadata)
                        declaration_generated = True
                        _add_section_break(doc)
                        suppress_next_page_break = True
                    skip_auto_block_content = True
                    first_heading1 = False
                    i += 1
                    continue

                # 特殊区块标题
                if '参考文献' in title_clean or 'references' in title_clean_lower:
                    in_abstract = False
                    in_references = True
                    if suppress_next_page_break:
                        suppress_next_page_break = False
                    else:
                        _ensure_page_break(doc)
                    para = doc.add_paragraph(_get_display_title(spec, '参考文献'), style='Heading 1')
                    first_heading1 = False
                    i += 1
                    continue
                
                if ('摘要' in title_clean or 'abstract' in title_clean_lower) and len(title_clean) < 15:
                    in_abstract = True
                    in_references = False
                    if suppress_next_page_break:
                        suppress_next_page_break = False
                    else:
                        _ensure_page_break(doc)
                    display_title = _get_display_title(spec, title_clean)
                    para = doc.add_paragraph()
                    title_style = _get_style(doc, 'Title')
                    if title_style is not None:
                        para.style = title_style
                    run = para.add_run(display_title)
                    run.font.name = spec.abstract.label_font if '摘要' in title_clean else spec.fonts.english_font
                    run.font.size = _resolve_front_title_size(spec)
                    run.font.bold = True
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), spec.abstract.label_font if '摘要' in title_clean else spec.fonts.english_font)
                    first_heading1 = False
                    i += 1
                    continue

                if ('目录' in title_clean or 'contents' in title_clean_lower) and len(title_clean) < 15:
                    in_abstract = False
                    in_references = False
                    if suppress_next_page_break:
                        suppress_next_page_break = False
                    else:
                        _ensure_page_break(doc)
                    para = doc.add_paragraph()
                    toc_style = _get_style(doc, 'TOC Heading')
                    if toc_style is not None:
                        para.style = toc_style
                    run = para.add_run(_get_display_title(spec, '目录'))
                    run.font.name = spec.fonts.heading_font
                    run.font.size = _resolve_front_title_size(spec)
                    run.font.bold = True
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), spec.fonts.heading_font)
                    # 插入真正的 TOC 域
                    _insert_toc_field(doc)
                    first_heading1 = False
                    i += 1
                    continue

                # ── 特殊章节标题（前言/结论/致谢/附录/文献综述）──
                # H1 级别但不编号，居中，分页
                _SPECIAL_CHAPTER_KEYWORDS = ("前言", "结论", "致谢", "附录", "文献综述",
                                             "foreword", "conclusion", "acknowledgement")
                is_special_chapter = any(k in title_clean.lower() for k in _SPECIAL_CHAPTER_KEYWORDS)
                if is_special_chapter and len(level_marks) <= 2:
                    in_abstract = False
                    in_references = False
                    if not body_section_started:
                        _add_section_break(doc)
                        body_section_started = True
                        suppress_next_page_break = True
                    if suppress_next_page_break:
                        suppress_next_page_break = False
                    else:
                        _ensure_page_break(doc)
                    para = doc.add_paragraph(style='Heading 1')
                    para.paragraph_format.space_after = Pt(12)
                    run = para.add_run(_get_display_title(spec, title_clean))
                    run.font.name = spec.fonts.heading_font
                    run.font.size = Pt(spec.font_sizes.heading_1)
                    run.font.bold = True
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), spec.fonts.heading_font)
                    first_heading1 = False
                    i += 1
                    continue
                
                # 论文题目（第一个一级标题，且不是"摘要"、"目录"等特殊标题，也不是章节标题）
                if first_heading1 and title_clean and len(level_marks) <= 2:
                    # 检查是否是特殊标题（摘要、目录等），这些不应该作为论文题目
                    special_titles = ['摘要', 'abstract', '目录', 'contents', '参考文献', 'references']
                    normalized_title = _normalize_title_token(title_clean)
                    is_special = normalized_title in {_normalize_title_token(item) for item in special_titles}
                    # 检查是否已经是章节标题（如"第一章 绪论"、"第1章 绪论"），如果是则不应作为论文题目
                    is_chapter_heading = bool(re.match(r'^第[一二三四五六七八九十百\d]+章', raw_title))

                    if not is_special and not is_chapter_heading:
                        in_abstract = False
                        # 使用论文标题样式（黑体小二）
                        para = doc.add_paragraph(title_clean, style='论文标题')
                        first_heading1 = False
                        i += 1
                        continue
                
                # 正规编号的标题
                in_abstract = False
                level = len(level_marks)
                # 兼容使用者错误的一阶标注习惯
                if level == 1 or (level == 2 and raw_title.startswith('第一章')):
                    new_body_section = False
                    if not body_section_started:
                        _add_section_break(doc)
                        body_section_started = True
                        new_body_section = True
                    # ── 前言/正文分节 ──
                    if in_front_matter and spec.use_roman_front_matter:
                        _add_section_break(doc)
                        in_front_matter = False

                    chap_num += 1
                    sec_num = 0
                    subsec_num = 0
                    subsubsec_num = 0
                    fig_num = 0
                    tab_num = 0
                    eq_num = 0
                    num_style = spec.heading_numbering.style
                    if num_style == "arabic":
                        heading_text = f"第{chap_num}章 {title_clean}"
                    else:
                        heading_text = f"第一章 {title_clean}" if chap_num == 1 else f"第{int_to_chinese(chap_num)}章 {title_clean}"
                    para = doc.add_paragraph(heading_text, style='Heading 1')
                    para.paragraph_format.page_break_before = not new_body_section
                elif level == 2:
                    sec_num += 1
                    subsec_num = 0
                    subsubsec_num = 0
                    if num_style == "arabic":
                        heading_text = f"{sec_num}. {title_clean}"
                    else:
                        heading_text = f"{int_to_chinese(sec_num)}、{title_clean}"
                    doc.add_paragraph(heading_text, style='Heading 2')
                elif level == 3:
                    subsec_num += 1
                    subsubsec_num = 0
                    if num_style == "arabic":
                        heading_text = f"{sec_num}.{subsec_num} {title_clean}"
                    else:
                        heading_text = f"（{int_to_chinese(subsec_num)}）{title_clean}"
                    doc.add_paragraph(heading_text, style='Heading 3')
                elif level >= 4:
                    subsubsec_num += 1
                    if num_style == "arabic":
                        heading_text = f"{sec_num}.{subsec_num}.{subsubsec_num} {title_clean}"
                    else:
                        heading_text = f"{subsubsec_num}. {title_clean}"
                    doc.add_paragraph(heading_text, style='Heading 4')
                
                first_heading1 = False
                just_parsed_table_caption = False
                i += 1
                continue
        
        # 参考文献条目（支持 [1]、[ 1 ] 等格式）
        if re.match(r'^\s*\[\s*\d+\s*\]', stripped):
            ref_text = stripped.strip()
            para = doc.add_paragraph()
            hang = spec.references.hanging_indent
            para.paragraph_format.left_indent = Cm(hang)
            para.paragraph_format.first_line_indent = Cm(-hang)
            # 使用新的格式化函数，确保数字和英文使用指定英文字体
            format_reference_text(para, ref_text, spec)
            i += 1
            continue
        
        # 表格处理（使用三线表格式）
        if stripped.startswith('|'):
            # 若之前没有出现专门的表题标注，则自动补充一个缺失的名字
            if i == 0 or not just_parsed_table_caption:
                tab_num += 1
                caption_text = f"表{chap_num}-{tab_num} （请补充表名）" if chap_num > 0 else f"表{tab_num} （请补充表名）"
                doc.add_paragraph(caption_text, style='Caption')
            
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            if len(table_lines) > 2:
                header_cells = [c.strip() for c in table_lines[0].split('|') if c.strip()]
                num_cols = len(header_cells)
                table = doc.add_table(rows=1, cols=num_cols)
                # 使用 Light Grid 样式作为基础，然后自定义边框
                table.style = 'Light Grid Accent 1'
                # 表格整体居中
                tblPr = table._tbl.tblPr
                if tblPr is None:
                    tblPr = OxmlElement('w:tblPr')
                    table._tbl.insert(0, tblPr)
                jc = OxmlElement('w:jc')
                jc.set(qn('w:val'), 'center')
                tblPr.append(jc)
                
                # 填充表头
                for col_idx, cell_text in enumerate(header_cells):
                    cell = table.cell(0, col_idx)
                    cell.text = cell_text
                    # 设置表头字体和居中
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            run.font.name = spec.fonts.english_font
                            run.font.size = Pt(spec.font_sizes.caption)
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), spec.fonts.chinese_font)
                            run.font.bold = True

                # 填充数据行
                for row_line in table_lines[2:]:
                    row_cells = [c.strip() for c in row_line.split('|')]
                    if row_line.startswith('|') and len(row_cells) > 0: row_cells.pop(0)
                    if row_line.endswith('|') and len(row_cells) > 0: row_cells.pop()
                    row = table.add_row()
                    for col_idx in range(min(num_cols, len(row_cells))):
                        cell = row.cells[col_idx]
                        cell.text = row_cells[col_idx]
                        # 设置单元格字体和居中
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in paragraph.runs:
                                run.font.name = spec.fonts.english_font
                                run.font.size = Pt(spec.font_sizes.caption)
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), spec.fonts.chinese_font)
                
                # 应用三线表格式
                set_table_borders(table)
                        
            just_parsed_table_caption = False
            continue
        
        # 分隔线
        if stripped.startswith('---'):
            i += 1
            continue
        
        # 引用块
        if stripped.startswith('>'):
            quote_text = stripped[1:].strip()
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Cm(1)
            run = para.add_run(quote_text)
            run.font.name = spec.fonts.abstract_font
            run.font.size = Pt(spec.font_sizes.body)
            run.font.italic = True
            run._element.rPr.rFonts.set(qn('w:eastAsia'), spec.fonts.abstract_font)
            i += 1
            continue
        
        # 列表项与无标签四级标题内联文本 (1. xxx: ...)
        if re.match(r'^([\*\-\+]|\d+\.)\s', stripped):
            list_text = re.sub(r'^([\*\-\+]|\d+\.)\s', '', stripped)
            prefix = re.match(r'^([\*\-\+]|\d+\.)\s', stripped).group(1)
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.left_indent = Pt(24)
            
            # 支持 “1. 标题：内容” 的行内黑体样式
            inline_title_match = re.match(r'^([^：:]+[：:])(.*)$', list_text)
            if prefix.endswith('.'):
                text_prefix = f"{prefix} "
            else:
                text_prefix = "• "
                
            if inline_title_match:
                run_title = para.add_run(text_prefix + inline_title_match.group(1))
                run_title.font.name = spec.fonts.heading_font
                run_title.font.size = Pt(spec.font_sizes.body)
                run_title.font.bold = True
                run_title._element.rPr.rFonts.set(qn('w:eastAsia'), spec.fonts.heading_font)

                run_content = para.add_run(inline_title_match.group(2))
                run_content.font.name = spec.fonts.chinese_font
                run_content.font.size = Pt(spec.font_sizes.body)
                run_content._element.rPr.rFonts.set(qn('w:eastAsia'), spec.fonts.chinese_font)
            else:
                run = para.add_run(text_prefix + list_text)
                run.font.name = spec.fonts.chinese_font
                run.font.size = Pt(spec.font_sizes.body)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), spec.fonts.chinese_font)
                
            just_parsed_table_caption = False
            i += 1
            continue
        
        # 正文段落
        if stripped and not stripped.startswith('#') and not stripped.startswith('['):
            abstract_line_match = re.match(
                r'^\s*(?:(\*\*|__))?\s*(摘要|Abstract)\s*(?:(\*\*|__))?\s*[:：]\s*(.*)$',
                stripped,
                re.IGNORECASE,
            )
            if abstract_line_match:
                in_abstract = True
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.paragraph_format.line_spacing = _resolve_body_line_spacing(spec)
                para.paragraph_format.first_line_indent = Cm(spec.abstract.first_line_indent)
                content = abstract_line_match.group(4).strip()
                run = para.add_run(content)
                run.font.size = Pt(spec.font_sizes.body)
                run.font.name = spec.fonts.english_font
                run._element.rPr.rFonts.set(qn('w:eastAsia'), spec.fonts.abstract_font)

                just_parsed_table_caption = False
                i += 1
                continue

            keyword_line_match = re.match(
                r'^\s*(?:(\*\*|__))?\s*(关键词|Keywords|Key words)\s*(?:(\*\*|__))?\s*[:：]\s*(.*)$',
                stripped,
                re.IGNORECASE,
            )
            if keyword_line_match:
                raw_label = keyword_line_match.group(2)
                content = keyword_line_match.group(4).strip()
                is_chinese_keywords = _normalize_title_token(raw_label) == _normalize_title_token("关键词")
                label_text = spec.abstract.keywords_label if is_chinese_keywords else spec.abstract.english_keywords_label
                para = doc.add_paragraph()
                run_label = para.add_run(label_text)
                run_label.font.name = spec.abstract.label_font if is_chinese_keywords else spec.fonts.english_font
                run_label.font.size = Pt(spec.font_sizes.body)
                run_label.font.bold = True
                run_label._element.rPr.rFonts.set(
                    qn('w:eastAsia'),
                    spec.abstract.label_font if is_chinese_keywords else spec.fonts.english_font,
                )

                run_rest = para.add_run(f"：{content}" if content else "：")
                run_rest.font.size = Pt(spec.font_sizes.body)
                run_rest.font.name = spec.fonts.english_font
                run_rest._element.rPr.rFonts.set(
                    qn('w:eastAsia'),
                    spec.fonts.chinese_font if is_chinese_keywords else spec.fonts.english_font,
                )

                just_parsed_table_caption = False
                i += 1
                continue

            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.line_spacing = _resolve_body_line_spacing(spec)
            para.paragraph_format.first_line_indent = Cm(spec.abstract.first_line_indent) if in_abstract else Pt(spec.font_sizes.body * 2)

            # 判断是否为摘要内容
            is_abstract_content = False
            if in_abstract and not any(
                _normalize_title_token(stripped).startswith(_normalize_title_token(label))
                for label in (spec.abstract.keywords_label, spec.abstract.english_keywords_label, "Keywords", "Key words")
            ):
                is_abstract_content = True

            # 处理正文中的行内公式和引用上标
            parts = re.split(r'(\$[^$]+\$|\[\d+(?:,\s*\d+)*\])', stripped)
            for part in parts:
                if not part:
                    continue
                # 行内公式 $...$
                if part.startswith('$') and part.endswith('$') and len(part) > 2:
                    _add_inline_equation(para, part[1:-1])
                # 引用上标 [n]
                elif re.match(r'^\[\d+(?:,\s*\d+)*\]$', part):
                    run = para.add_run(part)
                    run.font.size = Pt(spec.font_sizes.body)
                    run.font.name = spec.fonts.english_font
                    run.font.superscript = True
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), spec.fonts.abstract_font if is_abstract_content else spec.fonts.chinese_font)
                else:
                    run = para.add_run(part)
                    run.font.size = Pt(spec.font_sizes.body)
                    run.font.name = spec.fonts.english_font
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), spec.fonts.abstract_font if is_abstract_content else spec.fonts.chinese_font)
            just_parsed_table_caption = False
            i += 1
            continue
        
        i += 1

    # ── 后处理：设置页眉 + 页码 ──
    _apply_page_layout_to_sections(doc, spec)
    sections = doc.sections
    skip_header_sections = int(cover_generated) + int(declaration_generated)
    body_start_index = None
    if body_section_started:
        body_start_index = skip_header_sections + 1

    for idx, sec in enumerate(sections):
        if idx < skip_header_sections:
            _clear_section_header_footer(sec)
        else:
            add_bottom_border = not (idx == skip_header_sections and body_start_index is not None)
            _add_header(sec, spec, add_bottom_border=add_bottom_border)

    if body_start_index is not None and body_start_index < len(sections):
        for idx in range(skip_header_sections, body_start_index):
            sec = sections[idx]
            footer = sec.footer
            footer.is_linked_to_previous = False
            for paragraph in footer.paragraphs:
                paragraph.clear()
        body_sec = sections[body_start_index]
        _set_page_number_arabic(body_sec)
        _add_page_number_to_footer(body_sec, spec)
    elif len(sections) >= 2 and spec.use_roman_front_matter:
        _set_page_number_roman(sections[0])
        _add_page_number_to_footer(sections[0], spec)
        _set_page_number_arabic(sections[1])
        _add_page_number_to_footer(sections[1], spec)
    else:
        sec = sections[-1]
        _set_page_number_arabic(sec)
        _add_page_number_to_footer(sec, spec)

    # 保存文档
    doc.save(output_path)
    print("[OK] 转换完成！")
    print(f"输出文件：{output_path}")
    print("\n样式说明：")
    print("   - 一级标题：应用 'Heading 1' 样式（黑体小三号，居中），自动分页")
    print("   - 二级标题：应用 'Heading 2' 样式（黑体四号，靠左）")
    print("   - 三级标题：应用 'Heading 3' 样式（黑体小四号，靠左）")
    if spec.spacing.body_line_spacing_rule == "exact":
        print(f"   - 正文：宋体小四号，固定值{spec.spacing.body_line_spacing}磅，首行缩进2字符")
    else:
        print(f"   - 正文：宋体小四号，{spec.spacing.body_line_spacing}倍行距，首行缩进2字符")
    print("   - 参考文献：悬挂缩进格式")
    print("\n提示：在 Word 中可以通过'样式'面板直接修改各级标题样式")


def main():
    parser = argparse.ArgumentParser(
        description='将 Markdown 论文转换为 Word 文档，使用标准 Heading 样式',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog='''
使用示例:
  python convert_to_word.py input.md
  python convert_to_word.py input.md -o output.docx
        '''
    )
    
    parser.add_argument('input', help='输入的 Markdown 文件路径')
    parser.add_argument('-o', '--output', help='输出的 Word 文件路径（默认：input.docx）')
    
    args = parser.parse_args()
    
    # 检查输入文件
    if not os.path.exists(args.input):
        print(f"[ERROR] 文件不存在 - {args.input}")
        sys.exit(1)
    
    # 确定输出路径
    if args.output:
        output_path = args.output
    else:
        base_name = os.path.splitext(args.input)[0]
        output_path = f"{base_name}.docx"
    
    # 执行转换
    convert_markdown_to_word(args.input, output_path)


if __name__ == "__main__":
    main()
