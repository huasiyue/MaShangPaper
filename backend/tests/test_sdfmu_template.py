#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from __future__ import annotations

import subprocess
from pathlib import Path

from docx import Document

from app.services.document_pipeline import cleanup_paths, normalize_word_input
from app.services.formatters.engine import ThesisFormatEngine
from app.services.formatters.schools.sdfmu import create_sdfmu_spec
from app.services.word_converter import convert_markdown_to_word


def test_sdfmu_missing_english_abstract_warning():
    doc = Document()
    doc.add_paragraph("摘  要")
    doc.add_paragraph("这是中文摘要内容。" * 30)
    doc.add_paragraph("关键词：关键词一；关键词二；关键词三")
    doc.add_paragraph("目  录")

    engine = ThesisFormatEngine(create_sdfmu_spec("thesis"))
    report = engine.process_document(doc, "sample.docx")

    assert any("英文摘要" in issue.description for issue in report.issues)


def test_sdfmu_missing_declaration_page_warning():
    doc = Document()
    spec = create_sdfmu_spec("thesis")
    for label, _ in spec.cover.fields:
        doc.add_paragraph(f"{label}：测试内容")
    doc.add_paragraph("摘  要")
    doc.add_paragraph("这是中文摘要内容。" * 30)
    doc.add_paragraph("关键词：关键词一；关键词二；关键词三")
    doc.add_paragraph("Abstract")
    doc.add_paragraph("This is the abstract content. " * 20)
    doc.add_paragraph("Key words: keyword1; keyword2; keyword3")
    doc.add_paragraph("目  录")

    engine = ThesisFormatEngine(spec)
    report = engine.process_document(doc, "sample.docx")

    assert any("原创性保证书" in issue.description for issue in report.issues)


def test_sdfmu_reference_foreign_and_recent_requirements():
    doc = Document()
    doc.add_paragraph("参考文献")
    for idx in range(15):
        doc.add_paragraph(f"[{idx + 1}] 张三. 论文题目[J]. 某期刊, 2018, 1(1): 1-10.")

    engine = ThesisFormatEngine(create_sdfmu_spec("thesis"))
    report = engine.process_document(doc, "sample.docx")

    descriptions = [issue.description for issue in report.issues]
    assert any("外文文献数量较少" in description for description in descriptions)
    assert any("近5年" in description for description in descriptions)


def test_normalize_word_input_converts_doc(monkeypatch, tmp_path):
    source = tmp_path / "sample.doc"
    source.write_bytes(b"legacy-doc")

    def fake_run(cmd: list[str], check: bool, capture_output: bool, text: bool):
        output_path = Path(cmd[-1])
        output_path.write_bytes(b"fake-docx")
        return subprocess.CompletedProcess(cmd, 0, "", "")

    monkeypatch.setattr("app.services.document_pipeline.subprocess.run", fake_run)

    normalized_path, cleanup_targets = normalize_word_input(source)
    try:
        assert normalized_path.suffix == ".docx"
        assert normalized_path.exists()
        assert cleanup_targets == [normalized_path]
    finally:
        cleanup_paths(cleanup_targets)


def test_markdown_heading_levels_map_to_word_styles(tmp_path):
    markdown_path = tmp_path / "sample.md"
    output_path = tmp_path / "sample.docx"
    markdown_path.write_text(
        "\n".join(
            [
                "# 封面",
                "题目：测试论文题目",
                "专业：计算机科学与技术",
                "年级、班级：2022级本科1班",
                "# 摘要",
                "这是摘要内容。" * 20,
                "关键词：关键词一；关键词二；关键词三",
                "# 目录",
                "# 绪论",
                "## 研究背景与意义",
                "### 国内外研究现状",
                "#### 存在的问题",
            ]
        ),
        encoding="utf-8",
    )

    convert_markdown_to_word(str(markdown_path), str(output_path), "sdfmu", "thesis")
    doc = Document(str(output_path))

    expected_styles = {
        "摘  要": "Title",
        "目  录": "TOC Heading",
        "第1章 绪论": "Heading 1",
        "1. 研究背景与意义": "Heading 2",
        "1.1 国内外研究现状": "Heading 3",
        "1.1.1 存在的问题": "Heading 4",
    }

    actual_styles = {
        para.text.strip(): para.style.name
        for para in doc.paragraphs
        if para.text.strip() in expected_styles
    }
    assert actual_styles == expected_styles
    assert any("测试论文题目" in para.text for para in doc.paragraphs)
