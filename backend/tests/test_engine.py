#!/usr/bin/env python3
# -*- coding: utf-8 -*-
import pytest
from docx import Document

from app.services.formatters.engine import ThesisFormatEngine
from app.services.formatters.report import FormatIssueLevel, FormatReport
from app.services.formatters.spec import (
    AbstractRules,
    CaptionRules,
    CoverPageRules,
    FontAssignments,
    FontSizes,
    HeadingNumbering,
    HeadingRule,
    PageLayout,
    ReferenceRules,
    SignatureRules,
    SpacingRules,
    ThesisFormatSpec,
)


def _make_spec(school_id: str = "test", header_text: str = "Test") -> ThesisFormatSpec:
    return ThesisFormatSpec(
        school_id=school_id,
        school_name="Test University",
        header_text=header_text,
        page_layout=PageLayout(
            top_margin=2.54, bottom_margin=2.54, left_margin=3.17,
            right_margin=3.17, gutter=0, header_distance=1.5, footer_distance=1.75,
        ),
        fonts=FontAssignments(
            chinese_font="宋体",
            english_font="Times New Roman",
            heading_font="黑体",
            abstract_font="楷体",
            signature_font="仿宋体",
        ),
        font_sizes=FontSizes(
            title=22, heading_1=15, heading_2=14, heading_3=12,
            body=12, abstract_label=12, signature=12, header_footer=10.5, caption=10.5,
        ),
        spacing=SpacingRules(
            title_line_spacing=1.0, heading_line_spacing=1.0,
            body_line_spacing=1.5, paragraph_spacing=12,
        ),
        headings=(
            HeadingRule(alignment="center", font="黑体", size=22, bold=True),
            HeadingRule(alignment="center", font="黑体", size=15, bold=True),
            HeadingRule(alignment="left", font="黑体", size=14, bold=True),
            HeadingRule(alignment="left", font="宋体", size=12, bold=False),
            HeadingRule(alignment="left", font="宋体", size=12, bold=False),
        ),
        heading_numbering=HeadingNumbering(
            style="chinese",
            level_1=r"^第[一二三四五六七八九十百]+章",
            level_2=r"^[一二三四五六七八九十百]+、",
            level_3=r"^（[一二三四五六七八九十百]+）",
            level_4=r"^\d+\.",
        ),
        abstract=AbstractRules(
            min_chars=100, label_font="黑体", font_size=12,
            keywords=("摘要", "abstract"),
            boundary_keywords=("关键词", "keywords", "目录"),
            first_line_indent=0.74,
        ),
        references=ReferenceRules(
            standard="GB7714-87", min_count=10, max_typical_count=50,
            numbering_id=99, title_font="黑体", title_size=16,
            section_end_keywords=("致谢", "附录"),
            hanging_indent=0.74,
        ),
        signature=SignatureRules(
            keywords=("年级专业", "学生姓名", "指导教师", "学院", "学号"),
            max_position=30, font="仿宋体", size=12, alignment="center",
            prefix_patterns=(r"^(学生|教师|专业|班级|学号)[:：]",),
        ),
        caption=CaptionRules(
            pattern=r"^#?\s*(图|表)\s*\d+[-\.]\d+",
            alignment="center", font_size=10.5,
            line_spacing=1.0, space_before=6, space_after=6,
        ),
        cover=CoverPageRules(
            university_name="Test University",
            thesis_type_label="本科毕业论文（设计）",
            fields=(
                ("学院", "________"),
                ("专业", "________"),
                ("学生姓名", "________"),
            ),
            university_font="华文行楷",
            university_size=26,
            thesis_type_font="宋体",
            thesis_type_size=26,
            title_font="宋体",
            title_size=22,
            field_label_font="宋体",
            field_label_size=15,
            underline_width=6.0,
        ),
        special_titles=("摘要", "abstract", "参考文献", "references", "目录", "contents", "前言", "结论", "致谢", "附录", "文献综述"),
        body_first_line_indent=0.74,
    )


class TestDetectHeadingLevel:
    def test_chinese_level_1(self):
        engine = ThesisFormatEngine(_make_spec())
        assert engine._detect_heading_level("第一章 绪论") == 1

    def test_chinese_level_2(self):
        engine = ThesisFormatEngine(_make_spec())
        assert engine._detect_heading_level("一、研究背景") == 2

    def test_chinese_level_3(self):
        engine = ThesisFormatEngine(_make_spec())
        assert engine._detect_heading_level("（一）子标题") == 3

    def test_special_chapter_as_h1(self):
        engine = ThesisFormatEngine(_make_spec())
        assert engine._detect_heading_level("前言") == 1
        assert engine._detect_heading_level("结论") == 1
        assert engine._detect_heading_level("致谢") == 1
        assert engine._detect_heading_level("文献综述") == 1

    def test_thesis_title_heuristic(self):
        engine = ThesisFormatEngine(_make_spec())
        assert engine._detect_heading_level("基于深度学习的图像分类研究") == 0

    def test_body_text(self):
        engine = ThesisFormatEngine(_make_spec())
        assert engine._detect_heading_level("这是一段正文内容，带有标点符号。") == -1


class TestCheckAbstract:
    def test_missing_abstract(self):
        doc = Document()
        doc.add_paragraph("第一章 绪论")
        engine = ThesisFormatEngine(_make_spec())
        report = FormatReport(filename="test.docx")
        engine._check_abstract(doc, report)
        issues = [i for i in report.issues if "未检测到中文摘要" in i.description]
        assert any(i.level == FormatIssueLevel.WARNING for i in issues)

    def test_abstract_font_warning(self):
        doc = Document()
        p = doc.add_paragraph("摘要")
        for r in p.runs:
            r.font.name = "宋体"
        doc.add_paragraph("这是一段超过一百字的摘要内容。" * 10)
        engine = ThesisFormatEngine(_make_spec())
        report = FormatReport(filename="test.docx")
        engine._check_abstract(doc, report)
        issues = [i for i in report.issues if i.location == "中文摘要"]
        assert any("字体" in i.description for i in issues)

    def test_abstract_too_short(self):
        doc = Document()
        p = doc.add_paragraph("摘要内容很短。")
        for r in p.runs:
            r.font.name = "黑体"
        engine = ThesisFormatEngine(_make_spec())
        report = FormatReport(filename="test.docx")
        engine._check_abstract(doc, report)
        issues = [i for i in report.issues if i.location == "中文摘要"]
        assert any("过短" in i.description for i in issues)


class TestCheckReferences:
    def test_missing_references(self):
        doc = Document()
        doc.add_paragraph("第一章 绪论")
        engine = ThesisFormatEngine(_make_spec())
        report = engine.process_document(doc, "test.docx")
        ref_issues = [i for i in report.issues if "参考文献章节" in i.description]
        assert any(i.level == FormatIssueLevel.WARNING for i in ref_issues)

    def test_reference_count(self):
        doc = Document()
        doc.add_paragraph("参考文献")
        for i in range(3):
            doc.add_paragraph(f"[{i+1}] 作者. 文章名[J]. 期刊, 2024, 1(1): 1-10.")
        engine = ThesisFormatEngine(_make_spec())
        report = engine.process_document(doc, "test.docx")
        count_issues = [i for i in report.issues if "参考文献数量" in i.location]
        assert any(i.level == FormatIssueLevel.WARNING for i in count_issues)


class TestCheckCoverPage:
    def test_missing_cover_fields(self):
        doc = Document()
        doc.add_paragraph("第一章 绪论")
        engine = ThesisFormatEngine(_make_spec())
        report = engine.process_document(doc, "test.docx")
        engine._check_cover_page(doc, report)
        issues = [i for i in report.issues if i.location == "封面页"]
        assert any("学院" in i.description for i in issues)

    def test_cover_fields_present(self):
        doc = Document()
        doc.add_paragraph("学院 计算机学院")
        doc.add_paragraph("专业 软件工程")
        doc.add_paragraph("学生姓名 张三")
        engine = ThesisFormatEngine(_make_spec())
        report = engine.process_document(doc, "test.docx")
        issues = [i for i in report.issues if i.location == "封面页"]
        assert len(issues) == 0


class TestFormatParagraph:
    def test_special_title_formatted_as_h1(self):
        doc = Document()
        para = doc.add_paragraph("前言")
        engine = ThesisFormatEngine(_make_spec())
        engine._format_paragraph(para, 0, doc.paragraphs, engine.process_document(doc, "test.docx"))
        assert para.alignment is not None
        # H1 rule is bold=True, size=15
        for run in para.runs:
            assert run.font.bold is True
            assert run.font.size.pt == 15.0
