#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
扬州大学毕业论文格式自动转换工具
功能：
1. 读取待修改论文Word文档
2. 一键应用扬州大学毕业论文格式规范（区分毕业论文与毕业设计报告）
3. 生成格式审查报告

作者：AI Assistant
日期：2026-04-11
"""

import os
import re
import sys
import argparse
from dataclasses import dataclass, field
from typing import List, Dict, Tuple, Optional
from enum import Enum
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


class FormatIssueLevel(Enum):
    """格式问题严重程度"""
    ERROR = "错误"      # 必须修正
    WARNING = "警告"    # 建议修正
    INFO = "提示"       # 仅供参考


@dataclass
class FormatIssue:
    """格式问题记录"""
    level: FormatIssueLevel
    location: str       # 问题位置（如：第3页-标题）
    description: str    # 问题描述
    suggestion: str     # 修改建议
    current_value: str = ""  # 当前值
    expected_value: str = "" # 期望值


@dataclass
class FormatReport:
    """格式审查报告"""
    filename: str
    total_issues: int = 0
    error_count: int = 0
    warning_count: int = 0
    info_count: int = 0
    issues: List[FormatIssue] = field(default_factory=list)
    summary: str = ""


class ThesisType(Enum):
    """论文类型"""
    THESIS = "毕业论文"           # 毕业论文（一级标题居中，摘要内容楷体）
    DESIGN_REPORT = "毕业设计报告"  # 毕业设计报告（一级标题靠左，摘要内容宋体）


class YZUThesisFormat:
    """
    扬州大学毕业论文格式规范定义
    基于《扬州大学本科生毕业设计（论文）格式及要求》
    """

    def __init__(self, thesis_type: ThesisType = ThesisType.THESIS):
        self.thesis_type = thesis_type

    # ========== 页面设置 ==========
    PAGE_SETTINGS = {
        'top_margin': 2.2,      # 上边距 2.2cm
        'bottom_margin': 2.2,   # 下边距 2.2cm
        'left_margin': 2.5,     # 左边距 2.5cm
        'right_margin': 2.0,    # 右边距 2cm
        'gutter': 0.5,          # 装订线 0.5cm
        'header_distance': 1.2, # 页眉 1.2cm
        'footer_distance': 1.5, # 页脚 1.5cm
    }

    # ========== 字体设置 ==========
    FONT_SETTINGS = {
        'chinese_font': '宋体',
        'english_font': 'Times New Roman',
        'abstract_font_thesis': '楷体',      # 毕业论文摘要内容用楷体
        'abstract_font_design': '宋体',      # 设计报告摘要内容用宋体
        'heading_font': '黑体',
        'signature_font': '仿宋体',          # 署名用仿宋体
    }

    # ========== 字号设置（磅值） ==========
    FONT_SIZES = {
        'title': Pt(18),            # 小二号 = 18pt
        'heading_1': Pt(15),        # 小三号 = 15pt
        'heading_2': Pt(14),        # 四号 = 14pt
        'heading_3': Pt(12),        # 小四号 = 12pt
        'body': Pt(12),             # 小四号 = 12pt
        'abstract_label': Pt(12),   # 小四号
        'signature': Pt(12),        # 小四号
        'header_footer': Pt(9),     # 小五号 = 9pt
    }

    # ========== 行距设置 ==========
    LINE_SPACING = {
        'title': 1.0,
        'heading': 1.0,
        'body': 1.5,        # 1.5倍行距
        'paragraph_spacing': Pt(12),  # 段间距（小四号）
    }

    def get_heading_format(self, level: int) -> Tuple[str, WD_ALIGN_PARAGRAPH, str, Pt]:
        """
        获取标题层级格式
        关键差异：毕业论文一级标题居中，毕业设计报告一级标题靠左
        """
        # 毕业论文一级标题居中，设计报告靠左
        if level == 1:
            alignment = WD_ALIGN_PARAGRAPH.CENTER if self.thesis_type == ThesisType.THESIS else WD_ALIGN_PARAGRAPH.LEFT
        else:
            alignment = WD_ALIGN_PARAGRAPH.LEFT

        format_map = {
            0: ('', WD_ALIGN_PARAGRAPH.CENTER, '黑体', Pt(18)),  # 论文题目
            1: ('', alignment, '黑体', Pt(15)),  # 一级标题
            2: ('', WD_ALIGN_PARAGRAPH.LEFT, '黑体', Pt(14)),  # 二级标题
            3: ('', WD_ALIGN_PARAGRAPH.LEFT, '黑体', Pt(12)),  # 三级标题
            4: ('', WD_ALIGN_PARAGRAPH.LEFT, '黑体', Pt(12)),  # 四级标题
        }
        return format_map.get(level, ('', WD_ALIGN_PARAGRAPH.LEFT, '黑体', Pt(12)))

    def get_abstract_font(self) -> str:
        """获取摘要内容字体"""
        if self.thesis_type == ThesisType.THESIS:
            return self.FONT_SETTINGS['abstract_font_thesis']  # 楷体
        else:
            return self.FONT_SETTINGS['abstract_font_design']  # 宋体

    # ========== 论文结构顺序 ==========
    THESIS_STRUCTURE = [
        '封面',
        '中文摘要',
        '英文摘要',
        '目录',
        '绪论',
        '正文',
        '结论',
        '致谢',
        '参考文献',
        '附录',
    ]

    # ========== 摘要要求 ==========
    ABSTRACT_REQUIREMENTS = {
        'chinese_length': 300,      # 中文摘要约300字
        'label_font': '黑体',
        'font_size': Pt(12),
    }


class ThesisFormatter:
    """论文格式转换器"""

    def __init__(self, input_path: str, output_path: str = None, thesis_type: ThesisType = ThesisType.THESIS):
        self.input_path = input_path
        self.thesis_type = thesis_type  # 先设置 thesis_type
        self.output_path = output_path or self._generate_output_path(input_path)
        self.doc = None
        self.report = FormatReport(filename=os.path.basename(input_path))
        self.format_spec = YZUThesisFormat(thesis_type)

    def _generate_output_path(self, input_path: str) -> str:
        """生成输出文件路径"""
        dirname = os.path.dirname(input_path)
        basename = os.path.basename(input_path)
        name, ext = os.path.splitext(basename)
        return os.path.join(dirname, f"{name}_formatted{ext}")

    def load_document(self) -> bool:
        """加载Word文档"""
        try:
            self.doc = Document(self.input_path)
            self._add_issue(FormatIssueLevel.INFO, "文档加载", f"成功加载文档（按{self.thesis_type.value}格式处理）", "")
            return True
        except Exception as e:
            self._add_issue(FormatIssueLevel.ERROR, "文档加载", f"加载失败: {str(e)}", "检查文件路径和格式")
            return False

    def _add_issue(self, level: FormatIssueLevel, location: str, description: str,
                   suggestion: str, current: str = "", expected: str = ""):
        """添加格式问题记录"""
        issue = FormatIssue(
            level=level,
            location=location,
            description=description,
            suggestion=suggestion,
            current_value=current,
            expected_value=expected
        )
        self.report.issues.append(issue)
        self.report.total_issues += 1
        if level == FormatIssueLevel.ERROR:
            self.report.error_count += 1
        elif level == FormatIssueLevel.WARNING:
            self.report.warning_count += 1
        else:
            self.report.info_count += 1

    def apply_page_setup(self):
        """应用页面设置"""
        if not self.doc:
            return

        sections = self.doc.sections
        for section in sections:
            # 页面边距
            section.top_margin = Cm(self.format_spec.PAGE_SETTINGS['top_margin'])
            section.bottom_margin = Cm(self.format_spec.PAGE_SETTINGS['bottom_margin'])
            section.left_margin = Cm(self.format_spec.PAGE_SETTINGS['left_margin'])
            section.right_margin = Cm(self.format_spec.PAGE_SETTINGS['right_margin'])
            section.gutter = Cm(self.format_spec.PAGE_SETTINGS['gutter'])
            section.header_distance = Cm(self.format_spec.PAGE_SETTINGS['header_distance'])
            section.footer_distance = Cm(self.format_spec.PAGE_SETTINGS['footer_distance'])

        self._add_issue(FormatIssueLevel.INFO, "页面设置", "已应用扬州大学页面格式", "")

    def apply_header_footer(self):
        """应用页眉页脚设置"""
        if not self.doc:
            return

        for section in self.doc.sections:
            # 页眉
            header = section.header
            header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            header_para.text = "扬州大学本科生毕业论文"
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            for run in header_para.runs:
                run.font.name = self.format_spec.FONT_SETTINGS['chinese_font']
                run.font.size = self.format_spec.FONT_SIZES['header_footer']
                run._element.rPr.rFonts.set(qn('w:eastAsia'), self.format_spec.FONT_SETTINGS['chinese_font'])

            # 页脚 - 页码
            footer = section.footer
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 添加页码字段
            self._add_page_number(footer_para)

        self._add_issue(FormatIssueLevel.INFO, "页眉页脚", "已添加页眉和页码", "")

    def _add_page_number(self, paragraph):
        """添加页码字段"""
        run = paragraph.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

        run.font.name = self.format_spec.FONT_SETTINGS['english_font']
        run.font.size = self.format_spec.FONT_SIZES['body']

    def detect_heading_level(self, text: str) -> int:
        """检测标题层级"""
        text = text.strip()

        # 一级标题：第一章
        if re.match(r'^第[一二三四五六七八九十百]+章', text): return 1
        # 二级标题：一、
        if re.match(r'^[一二三四五六七八九十百]+、', text): return 2
        # 三级标题：（一）
        if re.match(r'^（[一二三四五六七八九十百]+）', text): return 3
        # 四级标题：1.
        if re.match(r'^\d+\.', text): return 4

        # 论文题目（居中的大字号文本）
        if len(text) < 50 and not any(c in text for c in '。，；：""''') and text:
            return 0

        return -1  # 非标题

    def is_signature_paragraph(self, text: str, para_index: int) -> bool:
        """检测是否为署名段落（增加位置和强度判定，防止误判正文中的‘年级’等词）"""
        if para_index > 30: # 署名通常只在封面或前几页
            return False
            
        text = text.strip()
        signature_keywords = ['年级专业', '学生姓名', '指导教师', '学院', '学号', '届别']
        count = 0
        for k in signature_keywords:
            if k in text: count += 1
            
        # 如果包含两个及以上关键词，或者是以特定引导词开头
        if count >= 2: return True
        if re.match(r'^(学生|教师|专业|班级|学号)[:：]', text): return True
        
        return False

    def is_abstract_content(self, paragraph_index: int, paragraphs: List) -> bool:
        """更健壮地检测是否为摘要内容段落"""
        start_idx = -1
        # 找摘要开始点
        for i in range(min(50, len(paragraphs))):
            t = paragraphs[i].text.strip().replace(' ', '').lower()
            # 允许独立行 (只有“摘要”二字) 或 行内形式 (如“摘要：本研究...”)
            if ('摘要' in t or 'abstract' in t):
                if len(t) < 15 or t.startswith('摘要') or t.startswith('abstract'):
                    start_idx = i
                    break
        
        if start_idx == -1 or paragraph_index <= start_idx:
            return False
            
        # 找摘要之后第一个边界点（关键词、目录、章节标题）
        boundary_idx = len(paragraphs)
        for j in range(start_idx + 1, len(paragraphs)):
            t = paragraphs[j].text.strip().replace(' ', '').lower()
            if '关键词' in t or 'keywords' in t or '目录' in t or t.startswith('第一章') or t.startswith('一、'):
                boundary_idx = j
                break
        
        return start_idx < paragraph_index < boundary_idx

    def is_special_title(self, text: str) -> bool:
        """检测是否为特殊标题（摘要、参考文献等），这些标题应保持原有格式"""
        text_clean = text.strip().replace(' ', '').lower()
        special_titles = ['摘要', 'abstract', '参考文献', 'references', '目录', 'contents']
        # 检查是否是独立的特殊标题（长度小于20，避免误判包含这些词的正文）
        if len(text) < 20:
            for title in special_titles:
                if title in text_clean:
                    return True
        return False

    def format_paragraph(self, paragraph, heading_level: int = -1, para_index: int = 0, all_paragraphs: List = None):
        """格式化段落"""
        text = paragraph.text.strip()

        if heading_level == -1:
            heading_level = self.detect_heading_level(text)

        # 检测特殊标题（摘要、参考文献等），保持原有格式不覆盖
        if self.is_special_title(text):
            # 只调整段落格式，不覆盖字体设置
            paragraph.paragraph_format.line_spacing = self.format_spec.LINE_SPACING['body']
        # 检测署名
        elif self.is_signature_paragraph(text, para_index):
            self._format_signature(paragraph)
        # 检测题注（图/表）
        elif re.match(r'^#?\s*(图|表)\s*\d+[-\.]\d+', text):
             self._format_caption(paragraph)
        # 检测摘要内容（毕业论文用楷体）
        elif self.thesis_type == ThesisType.THESIS and all_paragraphs and self.is_abstract_content(para_index, all_paragraphs):
            self._format_abstract_content(paragraph)
        elif heading_level >= 0:
            # 标题格式
            self._format_heading(paragraph, heading_level)
        else:
            # 正文格式
            self._format_body(paragraph)

    def _format_heading(self, paragraph, level: int):
        """格式化标题"""
        _, alignment, font_name, font_size = self.format_spec.get_heading_format(level)

        paragraph.alignment = alignment
        paragraph.paragraph_format.space_before = self.format_spec.LINE_SPACING['paragraph_spacing']
        paragraph.paragraph_format.space_after = self.format_spec.LINE_SPACING['paragraph_spacing']
        paragraph.paragraph_format.line_spacing = self.format_spec.LINE_SPACING['heading']

        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = font_size
            run.font.bold = True
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

    def _format_signature(self, paragraph):
        """格式化署名"""
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.line_spacing = self.format_spec.LINE_SPACING['body']

        for run in paragraph.runs:
            run.font.name = self.format_spec.FONT_SETTINGS['signature_font']
            run.font.size = self.format_spec.FONT_SIZES['signature']
            run._element.rPr.rFonts.set(qn('w:eastAsia'), self.format_spec.FONT_SETTINGS['signature_font'])

    def _format_abstract_content(self, paragraph):
        """格式化摘要内容（毕业论文用楷体）"""
        paragraph.paragraph_format.line_spacing = self.format_spec.LINE_SPACING['body']
        paragraph.paragraph_format.first_line_indent = Cm(0.74)

        abstract_font = self.format_spec.get_abstract_font()
        for run in paragraph.runs:
            run.font.size = self.format_spec.FONT_SIZES['body']
            run.font.name = self.format_spec.FONT_SETTINGS['english_font']
            text = run.text
            if text:
                if any('\u4e00' <= char <= '\u9fff' for char in text):
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), abstract_font)

    def _format_body(self, paragraph):
        """格式化正文"""
        paragraph.paragraph_format.line_spacing = self.format_spec.LINE_SPACING['body']
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进2字符

        for run in paragraph.runs:
            run.font.size = self.format_spec.FONT_SIZES['body']
            run.font.name = self.format_spec.FONT_SETTINGS['english_font'] # Times New Roman
            text = run.text
            if text:
                if any('\u4e00' <= char <= '\u9fff' for char in text):
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), self.format_spec.FONT_SETTINGS['chinese_font'])

    def _format_caption(self, paragraph):
        """格式化题注"""
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(6)
        
        for run in paragraph.runs:
            run.font.size = Pt(10.5)
            run.font.name = self.format_spec.FONT_SETTINGS['english_font']
            run._element.rPr.rFonts.set(qn('w:eastAsia'), self.format_spec.FONT_SETTINGS['chinese_font'])
            
    def process_document(self):
        """处理整个文档"""
        if not self.doc:
            return False

        # 应用页面设置
        self.apply_page_setup()

        # 应用页眉页脚
        self.apply_header_footer()

        # 处理所有段落
        paragraphs = self.doc.paragraphs
        for i, para in enumerate(paragraphs):
            self.format_paragraph(para, para_index=i, all_paragraphs=paragraphs)

        # 检查摘要
        self._check_abstract()

        # 格式化参考文献
        self.format_references()

        # 检查参考文献
        self._check_references()

        # 检查署名
        self._check_signature()

        return True

    def _check_abstract(self):
        """检查摘要格式"""
        abstract_found = False
        for para in self.doc.paragraphs[:20]:  # 检查前20段
            text = para.text.strip().replace(' ', '')
            if ('摘要' in text or 'abstract' in text.lower()) and len(text) < 20:
                abstract_found = True

                # 检查"摘要"标签格式
                for run in para.runs:
                    if '摘要' in run.text:
                        if run.font.name != self.format_spec.ABSTRACT_REQUIREMENTS['label_font']:
                            self._add_issue(
                                FormatIssueLevel.WARNING,
                                "中文摘要",
                                f"'摘要'标签字体应为黑体",
                                "修改字体为黑体",
                                run.font.name or "未设置",
                                "黑体"
                            )

                # 检查摘要内容长度
                content = text.replace('摘要', '').replace('：', '').replace(':', '').strip()
                if len(content) > 0 and len(content) < 100:
                    self._add_issue(
                        FormatIssueLevel.WARNING,
                        "中文摘要",
                        f"摘要内容可能过短（{len(content)}字），建议约300字",
                        "补充摘要内容",
                        f"{len(content)}字",
                        "约300字"
                    )

        if not abstract_found:
            self._add_issue(FormatIssueLevel.WARNING, "文档结构", "未检测到中文摘要", "请检查文档结构")

    def _check_references(self):
        """检查参考文献格式"""
        ref_found = False
        ref_section_started = False
        ref_items = []

        for i, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            text_no_space = text.replace(' ', '')

            # 检测参考文献标题
            if ('参考文献' in text_no_space or 'References' in text) and len(text) < 20:
                ref_found = True
                ref_section_started = True

                # 检查"参考文献"标题格式
                for run in para.runs:
                    if run.font.name != self.format_spec.FONT_SETTINGS['heading_font']:
                        self._add_issue(
                            FormatIssueLevel.WARNING,
                            "参考文献标题",
                            "'参考文献'标题字体应为黑体",
                            "修改字体为黑体",
                            run.font.name or "未设置",
                            "黑体"
                        )
                continue

            # 收集参考文献条目进行分析
            if ref_section_started and text and len(text) > 10:
                ref_items.append((i, text))

        if not ref_found:
            self._add_issue(FormatIssueLevel.WARNING, "文档结构", "未检测到参考文献章节", "请添加参考文献")
            return

        # 分析参考文献格式
        self._analyze_reference_format(ref_items)

    def _analyze_reference_format(self, ref_items: List[Tuple[int, str]]):
        """分析参考文献格式是否符合GB7714-87标准"""
        if not ref_items:
            self._add_issue(FormatIssueLevel.WARNING, "参考文献", "参考文献章节无具体条目", "请添加参考文献条目")
            return

        # 检测参考文献类型
        journal_count = 0
        book_count = 0
        conference_count = 0
        other_count = 0
        issues = []

        for idx, (para_idx, text) in enumerate(ref_items[:20]):  # 检查前20条
            # 期刊论文检测：包含期刊名、年份、卷号、期数、页码
            if re.search(r'[\u4e00-\u9fa5]+.*\d{4}.*\d+.*\d+.*\d+', text) and ('卷' in text or '期' in text):
                journal_count += 1
                # 检查期刊格式是否完整
                if not re.search(r'\d{4}.*[\(\（]\d+[\)\）]', text):  # 年份后应有期号
                    issues.append(f"第{idx+1}条期刊文献可能缺少期号格式")

            # 图书检测：包含出版社、出版年
            elif re.search(r'(出版社|出版|Press|Publishing)', text, re.IGNORECASE) and re.search(r'\d{4}', text):
                book_count += 1

            # 会议论文检测：包含"会议"、"论文集"、"Proceedings"等
            elif re.search(r'(会议|论文集|Proceedings|Conference|Symposium)', text, re.IGNORECASE):
                conference_count += 1

            else:
                other_count += 1

        # 添加统计信息
        total = len(ref_items)
        self._add_issue(
            FormatIssueLevel.INFO,
            "参考文献统计",
            f"检测到参考文献共{total}条：期刊论文约{journal_count}条，图书约{book_count}条，会议论文约{conference_count}条，其他约{other_count}条",
            ""
        )

        # 报告格式问题
        for issue in issues[:5]:  # 最多报告5个问题
            self._add_issue(FormatIssueLevel.WARNING, "参考文献格式", issue, "请参考GB7714-87标准格式")

        # 检查参考文献数量（本科论文通常要求10-20篇）
        if total < 8:
            self._add_issue(
                FormatIssueLevel.WARNING,
                "参考文献数量",
                f"参考文献数量较少（{total}条），本科论文通常要求10-20篇",
                "建议补充相关文献",
                f"{total}条",
                "10-20条"
            )
        elif total > 50:
            self._add_issue(
                FormatIssueLevel.INFO,
                "参考文献数量",
                f"参考文献数量较多（{total}条），请确保质量",
                ""
            )

    def format_references(self):
        """格式化参考文献章节，将手打编号改为Word自动编号"""
        ref_section_started = False
        ref_items = []  # 收集参考文献条目

        for para in self.doc.paragraphs:
            text = para.text.strip()
            text_no_space = text.replace(' ', '')

            # 检测参考文献标题
            if ('参考文献' in text_no_space or 'References' in text) and len(text) < 20:
                ref_section_started = True

                # 格式化标题
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in para.runs:
                    run.font.name = self.format_spec.FONT_SETTINGS['heading_font']
                    run.font.size = self.format_spec.FONT_SIZES['heading_1']
                    run.font.bold = True
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), self.format_spec.FONT_SETTINGS['heading_font'])
                continue

            # 收集参考文献条目（去除手打编号）
            if ref_section_started and text and len(text) > 10:
                # 移除手打编号 [数字] 或 数字.
                clean_text = re.sub(r'^\[\d+\]\s*', '', text)  # 移除 [1], [2] 等
                clean_text = re.sub(r'^\d+[.．]\s*', '', clean_text)  # 移除 1., 2. 等
                ref_items.append((para, clean_text))

            # 检测到下一章节则停止
            if ref_section_started and text and len(text) < 10:
                if any(keyword in text for keyword in ['致谢', '附录', '结论']):
                    break

        # 如果有参考文献条目，先创建编号定义
        if ref_items:
            self._create_reference_numbering()
        
        # 应用自动编号到参考文献条目
        for idx, (para, clean_text) in enumerate(ref_items, 1):
            # 清除原有内容
            para.clear()
            
            # 分离中英文和数字，分别应用不同字体
            self._add_reference_text_with_mixed_fonts(para, clean_text)
            
            # 设置段落格式
            para.paragraph_format.line_spacing = self.format_spec.LINE_SPACING['body']
            para.paragraph_format.first_line_indent = Cm(-0.74)  # 悬挂缩进
            para.paragraph_format.left_indent = Cm(0.74)  # 左缩进配合悬挂缩进
            para.paragraph_format.space_after = Pt(3)  # 条目间小间距
            
            # 应用自动编号 [1], [2], [3]...
            self._apply_reference_numbering(para, idx)

    def _apply_reference_numbering(self, paragraph, number: int):
        """
        应用参考文献自定义编号格式 [1], [2], [3]...
        使用Word的自动编号列表功能，而非手动添加编号
        """
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        
        # 获取段落的 pPr 元素
        pPr = paragraph._p.get_or_add_pPr()
        
        # 创建编号属性，使用编号ID 1（需要在文档中预定义或创建）
        numPr = parse_xml(r'''
            <w:numPr {}>
                <w:ilvl w:val="0"/>
                <w:numId w:val="99"/>
            </w:numPr>
        '''.format(nsdecls('w')))
        
        # 移除旧的编号设置
        for child in list(pPr):
            if child.tag.endswith('numPr'):
                pPr.remove(child)
        
        pPr.append(numPr)
        
        # 设置段落缩进（悬挂缩进效果）
        paragraph.paragraph_format.first_line_indent = None
        paragraph.paragraph_format.left_indent = None

    def _add_reference_text_with_mixed_fonts(self, paragraph, text: str):
        """
        为参考文献添加文本，中英文和数字使用不同字体
        中文使用宋体，英文和数字使用 Times New Roman
        """
        import re

        # 按字符遍历，区分中文和非中文（英文/数字）
        current_text = ""
        is_current_chinese = None

        for char in text:
            # 判断是否为中文
            is_chinese = '\u4e00' <= char <= '\u9fff'

            if is_current_chinese is None:
                is_current_chinese = is_chinese
                current_text = char
            elif is_current_chinese == is_chinese:
                current_text += char
            else:
                # 输出当前run
                if current_text:
                    run = paragraph.add_run(current_text)
                    run.font.size = self.format_spec.FONT_SIZES['body']
                    if is_current_chinese:
                        # 中文用宋体
                        run.font.name = self.format_spec.FONT_SETTINGS['chinese_font']
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), self.format_spec.FONT_SETTINGS['chinese_font'])
                    else:
                        # 英文和数字用 Times New Roman
                        run.font.name = self.format_spec.FONT_SETTINGS['english_font']
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), self.format_spec.FONT_SETTINGS['english_font'])
                # 开始新run
                is_current_chinese = is_chinese
                current_text = char

        # 输出最后一个run
        if current_text:
            run = paragraph.add_run(current_text)
            run.font.size = self.format_spec.FONT_SIZES['body']
            if is_current_chinese:
                run.font.name = self.format_spec.FONT_SETTINGS['chinese_font']
                run._element.rPr.rFonts.set(qn('w:eastAsia'), self.format_spec.FONT_SETTINGS['chinese_font'])
            else:
                run.font.name = self.format_spec.FONT_SETTINGS['english_font']
                run._element.rPr.rFonts.set(qn('w:eastAsia'), self.format_spec.FONT_SETTINGS['english_font'])

    def _create_reference_numbering(self):
        """
        创建参考文献的自定义编号格式 [1], [2], [3]...
        在文档的 numbering.xml 中定义编号样式
        """
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        
        # 访问文档的 numbering 部分
        try:
            numbering_part = self.doc.part.numbering_part
        except:
            # 如果没有 numbering 部分，创建一个
            from docx.parts.numbering import NumberingPart
            numbering_part = NumberingPart.new()
            self.doc.part._package.parts.append(numbering_part)
        
        # 创建抽象编号定义（编号格式 [1], [2]...）
        abstract_num_xml = parse_xml(r'''
            <w:abstractNum {} w:abstractNumId="99">
                <w:multiLevelType w:val="singleLevel"/>
                <w:lvl w:ilvl="0">
                    <w:numFmt w:val="decimal"/>
                    <w:lvlText w:val="[%1]"/>
                    <w:lvlJc w:val="left"/>
                    <w:pPr>
                        <w:ind w:left="0" w:firstLine="0"/>
                    </w:pPr>
                    <w:rPr>
                        <w:rFonts w:ascii="Times New Roman" w:eastAsia="Times New Roman"/>
                    </w:rPr>
                </w:lvl>
            </w:abstractNum>
        '''.format(nsdecls('w')))
        
        # 创建编号实例，设置起始编号为1
        num_xml = parse_xml(r'''
            <w:num {} w:numId="99">
                <w:abstractNumId w:val="99"/>
                <w:lvlOverride w:ilvl="0">
                    <w:startOverride w:val="1"/>
                </w:lvlOverride>
            </w:num>
        '''.format(nsdecls('w')))
        
        # 添加到 numbering 部分
        numbering_element = numbering_part._element
        numbering_element.append(abstract_num_xml)
        numbering_element.append(num_xml)

    def _check_signature(self):
        """检查署名"""
        signature_found = False
        for idx, para in enumerate(self.doc.paragraphs[:10]):
            if self.is_signature_paragraph(para.text, idx):
                signature_found = True
                break

        if self.thesis_type == ThesisType.THESIS and not signature_found:
            self._add_issue(FormatIssueLevel.INFO, "文档结构", "未检测到署名（年级专业、学生姓名、指导教师）", "毕业论文需要包含署名")

    def save_document(self) -> bool:
        """保存文档"""
        try:
            self.doc.save(self.output_path)
            self._add_issue(FormatIssueLevel.INFO, "文件保存", f"已保存到: {self.output_path}", "")
            return True
        except Exception as e:
            self._add_issue(FormatIssueLevel.ERROR, "文件保存", f"保存失败: {str(e)}", "检查输出路径权限")
            return False

    def generate_report(self) -> str:
        """生成格式审查报告"""
        lines = []
        lines.append("=" * 60)
        lines.append("扬州大学毕业论文格式审查报告")
        lines.append("=" * 60)
        lines.append(f"文件名: {self.report.filename}")
        lines.append(f"处理类型: {self.thesis_type.value}")
        lines.append(f"检查时间: {self._get_current_time()}")
        lines.append("")

        # 统计信息
        lines.append("-" * 60)
        lines.append("统计信息")
        lines.append("-" * 60)
        lines.append(f"问题总数: {self.report.total_issues}")
        lines.append(f"  - 错误: {self.report.error_count}")
        lines.append(f"  - 警告: {self.report.warning_count}")
        lines.append(f"  - 提示: {self.report.info_count}")
        lines.append("")

        # 按级别分组显示问题
        for level in [FormatIssueLevel.ERROR, FormatIssueLevel.WARNING, FormatIssueLevel.INFO]:
            issues = [i for i in self.report.issues if i.level == level]
            if issues:
                lines.append("-" * 60)
                lines.append(f"{level.value} ({len(issues)}项)")
                lines.append("-" * 60)

                for i, issue in enumerate(issues, 1):
                    lines.append(f"\n[{i}] 位置: {issue.location}")
                    lines.append(f"    问题: {issue.description}")
                    if issue.current_value and issue.expected_value:
                        lines.append(f"    当前值: {issue.current_value}")
                        lines.append(f"    期望值: {issue.expected_value}")
                    lines.append(f"    建议: {issue.suggestion}")

        lines.append("")
        lines.append("=" * 60)
        lines.append("格式规范参考")
        lines.append("=" * 60)
        lines.append(self._get_format_reference())

        return "\n".join(lines)

    def _get_current_time(self) -> str:
        """获取当前时间"""
        from datetime import datetime
        return datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    def _get_format_reference(self) -> str:
        """获取格式规范参考"""
        abstract_font = "楷体" if self.thesis_type == ThesisType.THESIS else "宋体"
        heading1_pos = "居中" if self.thesis_type == ThesisType.THESIS else "靠左顶格"

        ref = f"""
【处理类型】{self.thesis_type.value}

【页面设置】
- 纸张: A4
- 上边距: 2.2cm, 下边距: 2.2cm
- 左边距: 2.5cm, 右边距: 2cm
- 装订线: 0.5cm
- 页眉: 1.2cm, 页脚: 1.5cm

【字体字号】
- 论文题目: 黑体小二号(18pt), 居中
- 一级标题: 黑体小三号(15pt), {heading1_pos}
- 二级标题: 黑体四号(14pt), 靠左顶格
- 三级标题: 黑体小四号(12pt), 靠左顶格
- 正文: 宋体小四号(12pt), 1.5倍行距
- 英文/数字: Times New Roman
- 摘要内容: {abstract_font}小四号
- 署名: 仿宋体小四号, 居中

【标题层级】
- 一级: 一、二、三... 或 第一章
- 二级: （一）（二）... 或 第一节
- 三级: 1. 2. 3.
- 四级: （1）（2）...

【摘要要求】
- 中文摘要约300字
- "摘要"二字用黑体小四号, 居中
- 内容用{abstract_font}小四号

【参考文献】
- 格式标准: GB7714-87《文后参考文献著录规则》
- 数量要求: 本科论文通常10-20篇
- 期刊格式: 作者.文章题目名[J].期刊名,年份,卷号(期数):页码
- 图书格式: 作者.书名[M].出版地:出版单位,年份:页码
- 会议格式: 作者.文章题目名[C].会议名(论文集),年份:页码
- 格式示例:
  [1] 张三,李四.基于深度学习的图像识别研究[J].计算机科学,2023,50(3):45-52
  [2] 王五.人工智能导论[M].北京:清华大学出版社,2022:120-135
"""
        return ref

    def save_report(self, report_path: str = None):
        """保存审查报告"""
        if report_path is None:
            dirname = os.path.dirname(self.input_path)
            basename = os.path.basename(self.input_path)
            name, _ = os.path.splitext(basename)
            type_suffix = "论文" if self.thesis_type == ThesisType.THESIS else "设计报告"
            report_path = os.path.join(dirname, f"{name}_{type_suffix}_格式审查报告.txt")

        report_content = self.generate_report()
        with open(report_path, 'w', encoding='utf-8') as f:
            f.write(report_content)

        print(f"格式审查报告已保存: {report_path}")
        return report_path


def main():
    """主函数"""
    parser = argparse.ArgumentParser(
        description='扬州大学毕业论文格式自动转换工具',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog='''
使用示例:
  python yzu_thesis_formatter.py input.docx
  python yzu_thesis_formatter.py input.docx design
        '''
    )
    
    parser.add_argument('input', nargs='?', help='输入的 Word 文件路径')
    parser.add_argument('type', nargs='?', default='thesis', help='文档类型: thesis (毕业论文) 或 design (毕业设计报告)')
    
    args = parser.parse_args()

    print("=" * 60)
    print("扬州大学毕业论文格式自动转换工具")
    print("=" * 60)
    print()

    # 获取输入文件路径
    if args.input:
        input_path = args.input
        # 检查是否有类型参数
        thesis_type = ThesisType.DESIGN_REPORT if args.type.lower() in ['design', 'report', '设计'] else ThesisType.THESIS
    else:
        input_path = input("请输入待格式化论文的Word文件路径: ").strip().strip('"')

        # 选择论文类型
        print("\n请选择文档类型:")
        print("1. 毕业论文 (一级标题居中, 摘要内容楷体)")
        print("2. 毕业设计报告 (一级标题靠左, 摘要内容宋体)")
        type_choice = input("请输入选项(1/2, 默认1): ").strip()
        thesis_type = ThesisType.DESIGN_REPORT if type_choice == '2' else ThesisType.THESIS

    if not os.path.exists(input_path):
        print(f"错误: 文件不存在 - {input_path}")
        sys.exit(1)

    if not input_path.endswith(('.docx', '.doc')):
        print("错误: 请提供Word文档(.docx或.doc)")
        sys.exit(1)

    # 创建格式化器
    formatter = ThesisFormatter(input_path, thesis_type=thesis_type)

    # 加载文档
    print(f"正在加载文档: {input_path}")
    if not formatter.load_document():
        print("文档加载失败!")
        sys.exit(1)

    # 处理文档
    print(f"正在应用扬州大学{formatter.thesis_type.value}格式...")
    formatter.process_document()

    # 保存文档
    if formatter.save_document():
        print(f"格式化完成! 输出文件: {formatter.output_path}")
    else:
        print("保存失败!")

    # 生成并保存审查报告
    print("\n正在生成格式审查报告...")
    report_path = formatter.save_report()

    # 显示报告摘要
    print("\n" + "=" * 60)
    print("审查报告摘要")
    print("=" * 60)
    print(f"处理类型: {formatter.thesis_type.value}")
    print(f"问题总数: {formatter.report.total_issues}")
    print(f"  - 错误: {formatter.report.error_count}")
    print(f"  - 警告: {formatter.report.warning_count}")
    print(f"  - 提示: {formatter.report.info_count}")

    if formatter.report.error_count > 0:
        print("\n注意: 发现格式错误，请查看详细报告并修正!")
    elif formatter.report.warning_count > 0:
        print("\n提示: 发现格式警告，建议查看报告并优化!")
    else:
        print("\n很好! 文档格式符合规范!")

    print(f"\n详细报告请查看: {report_path}")


if __name__ == "__main__":
    main()
