#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将 Markdown 格式的论文初稿转换为 Word 文档
使用标准 Word Heading 样式，方便后续在 Word 中直接修改样式
"""

import os
import re
import sys
import argparse
from pathlib import Path
from urllib.parse import urlparse
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_page_break(paragraph):
    """在段落前添加分页符"""
    run = paragraph.add_run()
    run._r.append(OxmlElement('w:br'))
    run._r[-1].set(qn('w:type'), 'page')


def parse_image_meta(alt_text: str):
    """支持 `caption|width=8cm|align=center` 语法。"""
    parts = [part.strip() for part in alt_text.split('|') if part.strip()]
    caption = parts[0] if parts else "（请在此补充图名）"
    width = Cm(12)
    alignment = WD_ALIGN_PARAGRAPH.CENTER

    for option in parts[1:]:
        if '=' not in option:
            continue

        key, value = [item.strip().lower() for item in option.split('=', 1)]
        if key in {'width', 'w'}:
            width_match = re.match(r'^(\d+(?:\.\d+)?)(cm|mm|in|pt)?$', value)
            if not width_match:
                continue

            number = float(width_match.group(1))
            unit = width_match.group(2) or 'cm'
            if unit == 'cm':
                width = Cm(number)
            elif unit == 'mm':
                width = Cm(number / 10)
            elif unit == 'in':
                width = Inches(number)
            elif unit == 'pt':
                width = Pt(number)
        elif key in {'align', 'a'}:
            if value in {'left', '左'}:
                alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif value in {'right', '右'}:
                alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                alignment = WD_ALIGN_PARAGRAPH.CENTER

    return caption, width, alignment


def set_table_borders(table, border_type='three_line'):
    """
    设置表格边框为三线表格式
    三线表：顶线、栏目线、底线
    """
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    
    # 获取表格的 tbl 元素
    tbl = table._tbl
    
    # 创建表格属性
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(r'<w:tblPr {}></w:tblPr>'.format(nsdecls('w')))
    
    # 清除所有边框
    tblBorders = parse_xml(r'''
        <w:tblBorders {}>
            <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>
            <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>
            <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>
            <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>
            <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>
            <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>
        </w:tblBorders>
    '''.format(nsdecls('w')))
    
    # 移除旧的边框设置
    for child in list(tblPr):
        if child.tag.endswith('tblBorders'):
            tblPr.remove(child)
    
    tblPr.append(tblBorders)
    
    # 如果没有 tblPr，添加它
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)
    
    # 设置顶线（第一行顶部）和栏目线（第一行底部）
    if len(table.rows) > 0:
        header_row = table.rows[0]
        for cell in header_row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = parse_xml(r'''
                <w:tcBorders {}>
                    <w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/>
                    <w:bottom w:val="single" w:sz="6" w:space="0" w:color="000000"/>
                    <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>
                    <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>
                </w:tcBorders>
            '''.format(nsdecls('w')))
            # 移除旧的单元格边框设置
            for child in list(tcPr):
                if child.tag.endswith('tcBorders'):
                    tcPr.remove(child)
            tcPr.append(tcBorders)
    
    # 设置底线（最后一行底部）
    if len(table.rows) > 1:
        bottom_row = table.rows[-1]
        for cell in bottom_row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = parse_xml(r'''
                <w:tcBorders {}>
                    <w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>
                    <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>
                    <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>
                </w:tcBorders>
            '''.format(nsdecls('w')))
            # 移除旧的单元格边框设置
            for child in list(tcPr):
                if child.tag.endswith('tcBorders'):
                    tcPr.remove(child)
            tcPr.append(tcBorders)
    
    # 中间行无边框
    for row_idx in range(1, len(table.rows) - 1):
        row = table.rows[row_idx]
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = parse_xml(r'''
                <w:tcBorders {}>
                    <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>
                    <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>
                    <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>
                    <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>
                </w:tcBorders>
            '''.format(nsdecls('w')))
            # 移除旧的单元格边框设置
            for child in list(tcPr):
                if child.tag.endswith('tcBorders'):
                    tcPr.remove(child)
            tcPr.append(tcBorders)


def format_reference_text(para, text):
    """
    格式化参考文献文本，确保数字和英文使用 Times New Roman 字体
    中文字符使用宋体
    """
    import re
    
    # 匹配模式：序号、作者、标题、期刊、年份、卷期、页码等
    # 首先处理序号 [数字]
    parts = re.split(r'(\[\d+\])', text)
    
    for part in parts:
        if not part:
            continue
            
        # 如果是序号 [数字]
        if re.match(r'^\[\d+\]$', part):
            run = para.add_run(part)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        else:
            # 处理剩余部分，区分中英文
            # 按字符遍历，判断每个字符是中文还是英文/数字
            current_run_text = ""
            is_current_ascii = None
            
            for char in part:
                # 判断字符类型
                is_ascii = ord(char) < 128 and char.isprintable() or char.isspace()
                
                if is_current_ascii is None:
                    is_current_ascii = is_ascii
                    current_run_text = char
                elif is_current_ascii == is_ascii:
                    current_run_text += char
                else:
                    # 输出当前run
                    if current_run_text:
                        run = para.add_run(current_run_text)
                        run.font.size = Pt(12)
                        if is_current_ascii:
                            # 英文/数字用 Times New Roman
                            run.font.name = 'Times New Roman'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                        else:
                            # 中文用宋体
                            run.font.name = '宋体'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    # 开始新run
                    is_current_ascii = is_ascii
                    current_run_text = char
            
            # 输出最后一个run
            if current_run_text:
                run = para.add_run(current_run_text)
                run.font.size = Pt(12)
                if is_current_ascii:
                    run.font.name = 'Times New Roman'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                else:
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def setup_styles(doc):
    """设置文档的标准样式"""
    # 设置正文样式
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 设置 Heading 1 样式（一级标题）
    heading1 = doc.styles['Heading 1']
    heading1.font.name = 'Times New Roman'
    heading1.font.size = Pt(15)
    heading1.font.bold = True
    heading1.font.color.rgb = RGBColor(0, 0, 0)
    heading1._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    heading1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading1.paragraph_format.space_before = Pt(12)
    heading1.paragraph_format.space_after = Pt(12)
    heading1.paragraph_format.line_spacing = 1.0
    
    # 设置 Heading 2 样式（二级标题）
    heading2 = doc.styles['Heading 2']
    heading2.font.name = 'Times New Roman'
    heading2.font.size = Pt(14)
    heading2.font.bold = True
    heading2.font.color.rgb = RGBColor(0, 0, 0)
    heading2._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    heading2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading2.paragraph_format.space_before = Pt(12)
    heading2.paragraph_format.space_after = Pt(6)
    heading2.paragraph_format.line_spacing = 1.0
    
    # 设置 Heading 3 样式（三级标题）
    heading3 = doc.styles['Heading 3']
    heading3.font.name = 'Times New Roman'
    heading3.font.size = Pt(12)
    heading3.font.bold = True
    heading3.font.color.rgb = RGBColor(0, 0, 0)
    heading3._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    heading3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading3.paragraph_format.space_before = Pt(6)
    heading3.paragraph_format.space_after = Pt(6)
    heading3.paragraph_format.line_spacing = 1.0
    
    # 设置论文标题样式（自定义样式）
    if '论文标题' not in doc.styles:
        title_style = doc.styles.add_style('论文标题', WD_STYLE_TYPE.PARAGRAPH)
    else:
        title_style = doc.styles['论文标题']
    title_style.font.name = 'Times New Roman'
    title_style.font.size = Pt(18)  # 小二号
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0, 0, 0)
    title_style._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_before = Pt(12)
    title_style.paragraph_format.space_after = Pt(12)
    title_style.paragraph_format.line_spacing = 1.0
    
    # 设置 Caption 样式（题注）
    if 'Caption' not in doc.styles:
        caption_style = doc.styles.add_style('Caption', WD_STYLE_TYPE.PARAGRAPH)
    else:
        caption_style = doc.styles['Caption']
    caption_style.font.name = 'Times New Roman'
    caption_style.font.size = Pt(10.5)
    caption_style.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
    caption_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    caption_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_style.paragraph_format.line_spacing = 1.0
    caption_style.paragraph_format.space_before = Pt(6)
    caption_style.paragraph_format.space_after = Pt(6)


def int_to_chinese(num):
    chinese_digits = ["", "一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]
    if num == 0: return "零"
    if num <= 10: return chinese_digits[num]
    elif num < 20: return "十" + chinese_digits[num - 10]
    else: return chinese_digits[num // 10] + "十" + chinese_digits[num % 10]

def clean_heading(text):
    text = re.sub(r'^#{1,4}\s*', '', text)
    text = re.sub(r'^[第一二三四五六七八九十百]+章\s*', '', text)
    text = re.sub(r'^第[一二三四五六七八九十百]+节\s*', '', text)
    text = re.sub(r'^[一二三四五六七八九十]+\s*[、\.]?\s*', '', text)
    text = re.sub(r'^（[一二三四五六七八九十]+）\s*', '', text)
    text = re.sub(r'^[\d.]+、?\s*', '', text)
    text = re.sub(r'^（\d+）\s*', '', text)
    return text.strip()

def convert_markdown_to_word(markdown_path: str, output_path: str):
    """将 Markdown 转换为 Word 文档，使用标准 Heading 样式"""
    
    # 读取 Markdown 文件
    with open(markdown_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # 创建 Word 文档
    doc = Document()
    
    # 设置页面
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.0)
        section.gutter = Cm(0.5)
    
    # 设置标准样式
    setup_styles(doc)
    
    # 解析 Markdown 内容
    lines = content.split('\n')
    in_code_block = False
    in_abstract = False
    in_references = False
    first_heading1 = True  # 标记是否是第一个一级标题（论文题目）
    chap_num = 0
    sec_num = 0
    subsec_num = 0
    subsubsec_num = 0
    fig_num = 0
    tab_num = 0
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        
        # 跳过空行
        if not line.strip():
            i += 1
            continue
        
        stripped = line.strip()
        
        # 代码块处理
        if stripped.startswith('```'):
            in_code_block = not in_code_block
            if not in_code_block:
                # 代码块结束，添加代码内容
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].strip().startswith('```'):
                    code_lines.append(lines[i])
                    i += 1
                if code_lines:
                    para = doc.add_paragraph()
                    para.paragraph_format.left_indent = Cm(1)
                    run = para.add_run('\n'.join(code_lines))
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                i += 1
                continue
            else:
                i += 1
                continue
        
        if in_code_block:
            i += 1
            continue
        
        # -- 图片插图题注 --
        if stripped.startswith('!['):
            img_match = re.match(r'^!\[(.*?)\]\((.*?)\)$', stripped)
            if img_match:
                caption, image_width, image_alignment = parse_image_meta(img_match.group(1).strip())
                image_target = img_match.group(2).strip()
                fig_num += 1
                caption_text = f"图 {chap_num}-{fig_num} {caption}" if chap_num > 0 else f"图 {fig_num} {caption}"
                
                para_img = doc.add_paragraph()
                para_img.alignment = image_alignment
                run_img = para_img.add_run()

                image_path = Path(image_target)
                if image_path.exists():
                    run_img.add_picture(str(image_path), width=image_width)
                else:
                    run_img = para_img.add_run(f"[插图: {image_target}]")
                    run_img.font.color.rgb = RGBColor(128, 128, 128)
                
                # 图题注 (图片下方)
                para = doc.add_paragraph()
                para.alignment = image_alignment
                run = para.add_run(caption_text)
                run.font.name = '宋体'
                run.font.size = Pt(10.5)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                
                just_parsed_caption = True
                i += 1
                continue

        # -- 图/表 显示题注 --
        cap_match = re.match(r'^#*\s*(\*\*?)?(图|表)\s*(?:\d+[-\.]\d+|\d+)?[:：\s]+\s*(.+?)(\*\*?)?$', stripped)
        if cap_match and '格式' not in stripped:
            ctype = cap_match.group(2)
            caption = cap_match.group(3).strip()
            if ctype == '图':
                fig_num += 1
                caption_text = f"图 {chap_num}-{fig_num} {caption}" if chap_num > 0 else f"图 {fig_num} {caption}"
            else:
                tab_num += 1
                caption_text = f"表 {chap_num}-{tab_num} {caption}" if chap_num > 0 else f"表 {tab_num} {caption}"
            
            para = doc.add_paragraph(caption_text, style='Caption')
            just_parsed_table_caption = (ctype == '表')
            i += 1
            continue

        # -- 各类结构标题 --
        if stripped.startswith('#'):
            h_match = re.match(r'^(#{1,4})\s+(.*)', stripped)
            if h_match:
                level_marks = h_match.group(1)
                raw_title = h_match.group(2).strip()
                title_clean = clean_heading(raw_title)
                
                # 特殊区块标题
                if '参考文献' in title_clean or 'References' in title_clean:
                    in_abstract = False
                    in_references = True
                    para = doc.add_paragraph('参考文献', style='Heading 1')
                    if not first_heading1: para.paragraph_format.page_break_before = True
                    first_heading1 = False
                    i += 1
                    continue
                
                if '摘要' in title_clean or 'Abstract' in title_clean.lower() and len(title_clean) < 15:
                    if '摘要' in title_clean: in_abstract = True
                    para = doc.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = para.add_run(title_clean)
                    run.font.name = '黑体'
                    run.font.size = Pt(12)
                    run.font.bold = True
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                    first_heading1 = False
                    i += 1
                    continue

                if '目录' in title_clean or 'Contents' in title_clean.lower() and len(title_clean) < 15:
                    in_abstract = False
                    para = doc.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = para.add_run('目录')
                    run.font.name = '黑体'
                    run.font.size = Pt(15)
                    run.font.bold = True
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                    first_heading1 = False
                    i += 1
                    continue
                
                # 论文题目（第一个一级标题，且不是"摘要"、"目录"等特殊标题）
                if first_heading1 and title_clean and len(level_marks) <= 2:
                    # 检查是否是特殊标题（摘要、目录等），这些不应该作为论文题目
                    special_titles = ['摘要', 'abstract', '目录', 'contents', '参考文献', 'references']
                    is_special = any(s in title_clean.lower() for s in special_titles)
                    
                    if not is_special:
                        in_abstract = False
                        # 使用论文标题样式（黑体小二）
                        para = doc.add_paragraph(title_clean, style='论文标题')
                        first_heading1 = False
                        i += 1
                        continue
                
                # 正规编号的标题
                in_abstract = False
                level = len(level_marks)
                # 兼容使用者错误的一阶标注习惯
                if level == 1 or (level == 2 and raw_title.startswith('第一章')):
                    chap_num += 1
                    sec_num = 0
                    subsec_num = 0
                    subsubsec_num = 0
                    fig_num = 0
                    tab_num = 0
                    heading_text = f"第一章 {title_clean}" if chap_num==1 else f"第{int_to_chinese(chap_num)}章 {title_clean}"
                    para = doc.add_paragraph(heading_text, style='Heading 1')
                    para.paragraph_format.page_break_before = True
                elif level == 2:
                    sec_num += 1
                    subsec_num = 0
                    subsubsec_num = 0
                    heading_text = f"{int_to_chinese(sec_num)}、{title_clean}"
                    doc.add_paragraph(heading_text, style='Heading 2')
                elif level == 3:
                    subsec_num += 1
                    subsubsec_num = 0
                    heading_text = f"（{int_to_chinese(subsec_num)}）{title_clean}"
                    doc.add_paragraph(heading_text, style='Heading 3')
                elif level >= 4:
                    subsubsec_num += 1
                    heading_text = f"{subsubsec_num}. {title_clean}"
                    para = doc.add_paragraph(heading_text, style='Heading 3')
                    para.runs[0].font.size = Pt(12)
                
                first_heading1 = False
                just_parsed_table_caption = False
                i += 1
                continue
        
        # 参考文献条目（支持 [1]、[ 1 ] 等格式）
        if re.match(r'^\s*\[\s*\d+\s*\]', stripped):
            ref_text = stripped.strip()
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Cm(0.74)
            para.paragraph_format.first_line_indent = Cm(-0.74)
            # 使用新的格式化函数，确保数字和英文使用 Times New Roman
            format_reference_text(para, ref_text)
            i += 1
            continue
        
        # 表格处理（使用三线表格式）
        if stripped.startswith('|'):
            # 若之前没有出现专门的表题标注，则自动补充一个缺失的名字
            if i == 0 or not just_parsed_table_caption:
                tab_num += 1
                caption_text = f"表 {chap_num}-{tab_num} （请补充表名）" if chap_num > 0 else f"表 {tab_num} （请补充表名）"
                doc.add_paragraph(caption_text, style='Caption')
            
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            if len(table_lines) > 2:
                header_cells = [c.strip() for c in table_lines[0].split('|') if c.strip()]
                num_cols = len(header_cells)
                table = doc.add_table(rows=1, cols=num_cols)
                # 使用 Light Grid 样式作为基础，然后自定义边框
                table.style = 'Light Grid Accent 1'
                
                # 填充表头
                for col_idx, cell_text in enumerate(header_cells):
                    cell = table.cell(0, col_idx)
                    cell.text = cell_text
                    # 设置表头字体
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(10.5)
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                            run.font.bold = True
                
                # 填充数据行
                for row_line in table_lines[2:]:
                    row_cells = [c.strip() for c in row_line.split('|')]
                    if row_line.startswith('|') and len(row_cells) > 0: row_cells.pop(0)
                    if row_line.endswith('|') and len(row_cells) > 0: row_cells.pop()
                    row = table.add_row()
                    for col_idx in range(min(num_cols, len(row_cells))):
                        cell = row.cells[col_idx]
                        cell.text = row_cells[col_idx]
                        # 设置单元格字体
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(10.5)
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                
                # 应用三线表格式
                set_table_borders(table)
                        
            just_parsed_table_caption = False
            continue
        
        # 分隔线
        if stripped.startswith('---'):
            i += 1
            continue
        
        # 引用块
        if stripped.startswith('>'):
            quote_text = stripped[1:].strip()
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Cm(1)
            run = para.add_run(quote_text)
            run.font.name = '楷体'
            run.font.size = Pt(12)
            run.font.italic = True
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
            i += 1
            continue
        
        # 列表项与无标签四级标题内联文本 (1. xxx: ...)
        if re.match(r'^([\*\-\+]|\d+\.)\s', stripped):
            list_text = re.sub(r'^([\*\-\+]|\d+\.)\s', '', stripped)
            prefix = re.match(r'^([\*\-\+]|\d+\.)\s', stripped).group(1)
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.left_indent = Pt(24)
            
            # 支持 “1. 标题：内容” 的行内黑体样式
            inline_title_match = re.match(r'^([^：:]+[：:])(.*)$', list_text)
            if prefix.endswith('.'):
                text_prefix = f"{prefix} "
            else:
                text_prefix = "• "
                
            if inline_title_match:
                run_title = para.add_run(text_prefix + inline_title_match.group(1))
                run_title.font.name = '黑体'
                run_title.font.size = Pt(12)
                run_title.font.bold = True
                run_title._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                
                run_content = para.add_run(inline_title_match.group(2))
                run_content.font.name = '宋体'
                run_content.font.size = Pt(12)
                run_content._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            else:
                run = para.add_run(text_prefix + list_text)
                run.font.name = '宋体'
                run.font.size = Pt(12)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                
            just_parsed_table_caption = False
            i += 1
            continue
        
        # 正文段落
        if stripped and not stripped.startswith('#') and not stripped.startswith('['):
            # 特殊检测：加粗版摘要或关键词（针对截图1的情况）
            if re.match(r'^\s*(\*\*|__)\s*(摘要|关键词)\s*(\*\*|__)\s*[:：]', stripped):
                is_keywords_line = '关键词' in stripped
                if not is_keywords_line:
                    in_abstract = True
                para = doc.add_paragraph()
                if not is_keywords_line: para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                parts = re.split(r'([：:])', stripped, 1) # 分割前缀
                # 前缀（摘要/关键词）
                run_label = para.add_run(parts[0].replace('*','').replace('_',''))
                run_label.font.name = '黑体'
                run_label.font.size = Pt(12)
                run_label.font.bold = True
                run_label._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                
                # 冒号及之后内容
                if len(parts) > 1:
                    run_rest = para.add_run(''.join(parts[1:]))
                    run_rest.font.size = Pt(12)
                    if not is_keywords_line and in_abstract:
                        run_rest.font.name = 'Times New Roman'
                        run_rest._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
                    else:
                        run_rest.font.name = 'Times New Roman'
                        run_rest._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体' if is_keywords_line else '楷体')
                
                just_parsed_table_caption = False
                i += 1
                continue

            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.line_spacing = 1.5
            para.paragraph_format.first_line_indent = Pt(24)
            
            # 判断是否为摘要内容
            is_abstract_content = False
            if in_abstract and '关键词' not in stripped:
                is_abstract_content = True
            
            # 处理正文中的引用上标
            parts = re.split(r'(\[\d+(?:,\s*\d+)*\])', stripped)
            for part in parts:
                run = para.add_run(part)
                run.font.size = Pt(12)
                run.font.name = 'Times New Roman' # 强制英文罗马体
                if re.match(r'^\[\d+(?:,\s*\d+)*\]$', part):
                    run.font.superscript = True
                if is_abstract_content:
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
                else:
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            just_parsed_table_caption = False
            i += 1
            continue
        
        i += 1
    
    # 保存文档
    doc.save(output_path)
    print("[OK] 转换完成！")
    print(f"输出文件：{output_path}")
    print("\n样式说明：")
    print("   - 一级标题：应用 'Heading 1' 样式（黑体小三号，居中），自动分页")
    print("   - 二级标题：应用 'Heading 2' 样式（黑体四号，靠左）")
    print("   - 三级标题：应用 'Heading 3' 样式（黑体小四号，靠左）")
    print("   - 正文：宋体小四号，1.5倍行距，首行缩进2字符")
    print("   - 参考文献：悬挂缩进格式")
    print("\n提示：在 Word 中可以通过'样式'面板直接修改各级标题样式")


def main():
    parser = argparse.ArgumentParser(
        description='将 Markdown 论文转换为 Word 文档，使用标准 Heading 样式',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog='''
使用示例:
  python convert_to_word.py input.md
  python convert_to_word.py input.md -o output.docx
        '''
    )
    
    parser.add_argument('input', help='输入的 Markdown 文件路径')
    parser.add_argument('-o', '--output', help='输出的 Word 文件路径（默认：input.docx）')
    
    args = parser.parse_args()
    
    # 检查输入文件
    if not os.path.exists(args.input):
        print(f"[ERROR] 文件不存在 - {args.input}")
        sys.exit(1)
    
    # 确定输出路径
    if args.output:
        output_path = args.output
    else:
        base_name = os.path.splitext(args.input)[0]
        output_path = f"{base_name}.docx"
    
    # 执行转换
    convert_markdown_to_word(args.input, output_path)


if __name__ == "__main__":
    main()
